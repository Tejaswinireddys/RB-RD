#!/usr/bin/env python3
"""
Generate RabbitMQ Failover Scenarios and Test Cases Documentation
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def add_code_block(doc, text):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.style = 'Intense Quote'
    return para

def add_table_with_header(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    return table

def add_warning_box(doc, text):
    """Add a warning/note box"""
    para = doc.add_paragraph()
    run = para.add_run('⚠️  WARNING: ' + text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(204, 102, 0)
    return para

def add_note_box(doc, text):
    """Add an info/note box"""
    para = doc.add_paragraph()
    run = para.add_run('📝 NOTE: ' + text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 102, 204)
    return para

def create_failover_document():
    """Create the RabbitMQ failover scenarios documentation"""
    doc = Document()

    # Title Page
    title = doc.add_heading('RabbitMQ Cluster Failover', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Scenarios, Test Cases, and Procedures', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    platform = doc.add_heading('Three-Node Cluster Configuration', 3)
    platform.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Test Environment Setup',
        '3. Prerequisites and Assumptions',
        '4. Cluster Architecture Overview',
        '5. Failover Scenarios Overview',
        '6. Test Case 1: Master Node Failure',
        '7. Test Case 2: Replica Node Failure',
        '8. Test Case 3: Network Partition (Split-Brain)',
        '9. Test Case 4: Graceful Node Shutdown',
        '10. Test Case 5: Multiple Node Failure',
        '11. Test Case 6: Disk Space Exhaustion',
        '12. Test Case 7: Memory Pressure',
        '13. Test Case 8: Rolling Restart',
        '14. Test Case 9: Queue Mirroring Validation',
        '15. Test Case 10: Client Connection Failover',
        '16. Monitoring and Validation Commands',
        '17. Common Issues and Troubleshooting',
        '18. Recovery Procedures',
        '19. Best Practices for Production',
        '20. Test Results Template'
    ]
    for item in toc_items:
        add_paragraph(doc, item)

    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc,
        'This document provides comprehensive failover scenarios and test cases for a three-node '
        'RabbitMQ cluster deployment. It includes detailed step-by-step procedures to validate '
        'high availability, automatic failover, queue mirroring, and disaster recovery capabilities.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Purpose of Failover Testing:', bold=True)
    purposes = [
        'Validate automatic failover mechanisms work as expected',
        'Ensure zero or minimal message loss during node failures',
        'Test cluster behavior under various failure conditions',
        'Verify client reconnection and message delivery continuity',
        'Identify potential weaknesses in the cluster configuration',
        'Document recovery procedures for production incidents',
        'Build confidence in the high availability setup'
    ]
    for purpose in purposes:
        doc.add_paragraph(purpose, style='List Bullet')

    doc.add_page_break()

    # 2. Test Environment Setup
    add_heading(doc, '2. Test Environment Setup', 1)

    add_paragraph(doc, 'Required Infrastructure:', bold=True)
    infra = [
        ['Component', 'Specification', 'Purpose'],
        ['3 RHEL 8 Servers', '4 CPU, 8GB RAM each', 'RabbitMQ cluster nodes'],
        ['Test Client Machine', '2 CPU, 4GB RAM', 'Producer/consumer testing'],
        ['Network Access', 'All nodes can communicate', 'Cluster formation'],
        ['Monitoring Tools', 'RabbitMQ Management UI', 'Cluster status monitoring']
    ]
    add_table_with_header(doc, infra[0], infra[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Cluster Node Details:', bold=True)
    nodes = [
        ['Hostname', 'IP Address', 'Role', 'Node Name'],
        ['rabbitmq-node1', '192.168.1.101', 'Master/Primary', 'rabbit@rabbitmq-node1'],
        ['rabbitmq-node2', '192.168.1.102', 'Replica', 'rabbit@rabbitmq-node2'],
        ['rabbitmq-node3', '192.168.1.103', 'Replica', 'rabbit@rabbitmq-node3']
    ]
    add_table_with_header(doc, nodes[0], nodes[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Software Versions:', bold=True)
    add_code_block(doc,
        'RabbitMQ Version: 4.1.x\n'
        'Erlang Version: 26.x\n'
        'Operating System: RHEL 8.x'
    )

    doc.add_page_break()

    # 3. Prerequisites and Assumptions
    add_heading(doc, '3. Prerequisites and Assumptions', 1)

    add_paragraph(doc, 'Prerequisites:', bold=True)
    prereqs = [
        'RabbitMQ cluster is fully operational with 3 nodes',
        'All nodes are properly clustered and synchronized',
        'Queue mirroring policy is configured (ha-all or ha-exactly:2)',
        'RabbitMQ Management plugin is enabled on all nodes',
        'Administrative access to all cluster nodes',
        'Monitoring tools are configured and accessible',
        'Test queues and exchanges are created',
        'Test client applications are ready'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Initial Cluster Verification:', bold=True)
    add_code_block(doc,
        '# Run on any node to verify cluster status\n'
        'sudo rabbitmqctl cluster_status\n\n'
        '# Expected output:\n'
        '# - All 3 nodes listed under "Running Nodes"\n'
        '# - All 3 nodes listed under "Disk Nodes"\n\n'
        '# Verify queue mirroring policy\n'
        'sudo rabbitmqctl list_policies\n\n'
        '# Check node health\n'
        'sudo rabbitmq-diagnostics check_running\n'
        'sudo rabbitmq-diagnostics check_local_alarms'
    )

    doc.add_page_break()

    # 4. Cluster Architecture Overview
    add_heading(doc, '4. Cluster Architecture Overview', 1)

    add_paragraph(doc, 'Three-Node Cluster Configuration:', bold=True)
    add_paragraph(doc,
        'In a three-node RabbitMQ cluster, all nodes are equal peers. However, for each queue, '
        'one node acts as the master (leader) and others as mirrors (replicas). The master node '
        'handles all operations for that queue, while mirrors maintain synchronized copies.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Key Concepts:', bold=True)
    concepts = [
        ('Queue Master', 'The node that hosts the queue and handles all read/write operations'),
        ('Queue Mirror', 'Replicas that maintain synchronized copies of the queue'),
        ('Quorum', 'Minimum number of nodes required for cluster decisions (majority: 2 out of 3)'),
        ('Synchronization', 'Process of replicating messages from master to mirrors'),
        ('Failover', 'Automatic promotion of a mirror to master when master fails'),
        ('Partition Handling', 'Strategy for dealing with network splits (autoheal/pause_minority)')
    ]
    for concept, description in concepts:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(concept + ': ').bold = True
        para.add_run(description)

    doc.add_paragraph()
    add_paragraph(doc, 'High Availability Policy:', bold=True)
    add_code_block(doc,
        '# View current HA policy\n'
        'sudo rabbitmqctl list_policies\n\n'
        '# Expected policy (mirror to all nodes):\n'
        'Name: ha-all\n'
        'Pattern: .*\n'
        'Definition: {"ha-mode":"all","ha-sync-mode":"automatic"}'
    )

    doc.add_page_break()

    # 5. Failover Scenarios Overview
    add_heading(doc, '5. Failover Scenarios Overview', 1)

    add_paragraph(doc, 'Test Scenarios Summary:', bold=True)
    scenarios = [
        ['Scenario', 'Type', 'Impact Level', 'Expected Recovery Time'],
        ['Master Node Failure', 'Hard Failure', 'Medium', '< 30 seconds'],
        ['Replica Node Failure', 'Hard Failure', 'Low', 'Immediate (no impact)'],
        ['Network Partition', 'Split-Brain', 'High', '1-2 minutes'],
        ['Graceful Shutdown', 'Planned', 'Low', 'Immediate (no impact)'],
        ['Multiple Node Failure', 'Catastrophic', 'Critical', 'Manual intervention'],
        ['Disk Space Exhaustion', 'Resource', 'High', 'After disk cleanup'],
        ['Memory Pressure', 'Resource', 'Medium', '< 1 minute'],
        ['Rolling Restart', 'Maintenance', 'None', 'N/A'],
        ['Queue Mirroring Test', 'Validation', 'None', 'N/A'],
        ['Client Failover', 'Application', 'Low', '< 10 seconds']
    ]
    add_table_with_header(doc, scenarios[0], scenarios[1:])

    doc.add_page_break()

    # 6. Test Case 1: Master Node Failure
    add_heading(doc, '6. Test Case 1: Master Node Failure', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Simulate a catastrophic failure of the node hosting the queue master and verify automatic '
        'failover to a mirror node with minimal message loss and downtime.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: HIGH', bold=True)
    add_paragraph(doc, 'Expected Outcome: Automatic failover within 30 seconds', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Pre-Test Setup:', bold=True)
    add_code_block(doc,
        '# 1. Create test queue\n'
        'sudo rabbitmqadmin declare queue name=test-failover-queue durable=true\n\n'
        '# 2. Verify queue master location\n'
        'sudo rabbitmqctl list_queues name pid slave_pids synchronised_slave_pids\n\n'
        '# 3. Start message producer (keep running)\n'
        'python3 producer.py --queue test-failover-queue --rate 100\n\n'
        '# 4. Start message consumer (keep running)\n'
        'python3 consumer.py --queue test-failover-queue'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    steps = [
        ('Step 1: Record Initial State',
         'sudo rabbitmqctl cluster_status\n'
         'sudo rabbitmqctl list_queues name messages consumers\n'
         '# Note the queue master node'),

        ('Step 2: Identify Master Node',
         'sudo rabbitmqctl list_queues name pid | grep test-failover-queue\n'
         '# The node name in the pid indicates the master'),

        ('Step 3: Simulate Hard Failure (Kill Master Node)',
         '# On the master node:\n'
         'sudo systemctl stop rabbitmq-server\n'
         '# OR force kill process:\n'
         'sudo killall -9 beam.smp'),

        ('Step 4: Monitor Failover (On Another Node)',
         'watch -n 1 "sudo rabbitmqctl cluster_status"\n'
         '# Observe master node marked as down\n'
         '# Watch for mirror promotion'),

        ('Step 5: Verify Queue Availability',
         'sudo rabbitmqctl list_queues name messages consumers state\n'
         '# Queue should still be running on mirror node'),

        ('Step 6: Check Client Connections',
         'sudo rabbitmqctl list_connections name state\n'
         '# Clients should reconnect to available nodes'),

        ('Step 7: Verify Message Flow',
         '# Check producer and consumer logs\n'
         '# Verify messages continue to flow\n'
         '# Note any message loss (should be minimal)'),

        ('Step 8: Check Management UI',
         'http://192.168.1.102:15672\n'
         '# Verify cluster shows 2 nodes running\n'
         '# Check queue is accessible and processing messages')
    ]

    for step_title, step_cmd in steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Expected Results:', bold=True)
    results = [
        'Master node is marked as down in cluster status',
        'One of the mirror nodes is promoted to master automatically',
        'Queue remains available and continues processing messages',
        'Clients reconnect to available nodes within 10-15 seconds',
        'Message loss is minimal or zero (only in-flight messages)',
        'No manual intervention required',
        'Cluster operates with 2 nodes until failed node is restored'
    ]
    for result in results:
        doc.add_paragraph(result, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Post-Test Recovery:', bold=True)
    add_code_block(doc,
        '# On the failed node:\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# Verify node rejoins cluster\n'
        'sudo rabbitmqctl cluster_status\n\n'
        '# Check queue synchronization\n'
        'sudo rabbitmqctl list_queues name synchronised_slave_pids'
    )

    add_warning_box(doc,
        'In-flight messages that were not acknowledged may be lost during hard failure. '
        'Use publisher confirms and consumer acknowledgments for critical messages.')

    doc.add_page_break()

    # 7. Test Case 2: Replica Node Failure
    add_heading(doc, '7. Test Case 2: Replica Node Failure', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Verify that failure of a replica (mirror) node has minimal impact on cluster operations '
        'and message processing continues normally.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: LOW', bold=True)
    add_paragraph(doc, 'Expected Outcome: No impact on service availability', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    replica_steps = [
        ('Step 1: Identify Replica Node',
         'sudo rabbitmqctl list_queues name pid slave_pids\n'
         '# Choose a node from slave_pids (not the master)'),

        ('Step 2: Stop Replica Node',
         '# On the replica node:\n'
         'sudo systemctl stop rabbitmq-server'),

        ('Step 3: Verify Cluster Status',
         '# On master or remaining replica:\n'
         'sudo rabbitmqctl cluster_status\n'
         '# Replica node should be marked as down'),

        ('Step 4: Verify Queue Operations',
         'sudo rabbitmqctl list_queues name messages consumers\n'
         '# Queue should continue normal operations'),

        ('Step 5: Test Message Flow',
         '# Producer and consumer should continue without interruption\n'
         '# No reconnection should be needed if not connected to failed node'),

        ('Step 6: Restart Failed Node',
         '# On the failed replica node:\n'
         'sudo systemctl start rabbitmq-server'),

        ('Step 7: Verify Synchronization',
         'sudo rabbitmqctl list_queues name synchronised_slave_pids\n'
         '# Wait for queue to resynchronize with the rejoined node')
    ]

    for step_title, step_cmd in replica_steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Expected Results:', bold=True)
    replica_results = [
        'No impact on message processing',
        'Queue master continues operating normally',
        'Clients connected to other nodes are unaffected',
        'Clients connected to failed replica reconnect to other nodes',
        'No message loss occurs',
        'Failed node rejoins cluster cleanly after restart',
        'Queue resynchronizes automatically with the rejoined node'
    ]
    for result in replica_results:
        doc.add_paragraph(result, style='List Bullet')

    add_note_box(doc,
        'Replica failure is the least impactful failure scenario. The cluster can lose up to '
        'n-1 replicas (where n is total nodes) and still maintain full functionality.')

    doc.add_page_break()

    # 8. Test Case 3: Network Partition (Split-Brain)
    add_heading(doc, '8. Test Case 3: Network Partition (Split-Brain)', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Test cluster behavior when network connectivity is lost between nodes, creating a '
        'partition scenario. Verify partition handling strategy (autoheal or pause_minority) '
        'works correctly.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: CRITICAL', bold=True)
    add_paragraph(doc, 'Expected Outcome: Partition detected and healed automatically', italic=True)

    doc.add_paragraph()
    add_warning_box(doc,
        'Network partitions are the most dangerous failure scenario. Improper handling can lead '
        'to data inconsistency and message loss.')

    doc.add_paragraph()
    add_paragraph(doc, 'Prerequisites:', bold=True)
    add_code_block(doc,
        '# Verify partition handling mode\n'
        'sudo rabbitmqctl eval "application:get_env(rabbit, cluster_partition_handling)."\n\n'
        '# Should be set to "autoheal" (recommended) or "pause_minority"\n'
        '# In /etc/rabbitmq/rabbitmq.conf:\n'
        'cluster_partition_handling = autoheal'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    partition_steps = [
        ('Step 1: Record Pre-Partition State',
         'sudo rabbitmqctl cluster_status\n'
         'sudo rabbitmqctl list_queues name messages\n'
         '# All 3 nodes should be running'),

        ('Step 2: Create Network Partition',
         '# Method 1: Using iptables (recommended for testing)\n'
         '# On node1, block traffic from node2 and node3:\n'
         'sudo iptables -A INPUT -s 192.168.1.102 -j DROP\n'
         'sudo iptables -A INPUT -s 192.168.1.103 -j DROP\n'
         'sudo iptables -A OUTPUT -d 192.168.1.102 -j DROP\n'
         'sudo iptables -A OUTPUT -d 192.168.1.103 -j DROP\n\n'
         '# This creates partition: [node1] | [node2, node3]'),

        ('Step 3: Detect Partition',
         '# On node1:\n'
         'sudo rabbitmqctl cluster_status\n'
         '# Should show node2 and node3 as partitioned\n\n'
         '# On node2 or node3:\n'
         'sudo rabbitmqctl cluster_status\n'
         '# Should show node1 as partitioned'),

        ('Step 4: Observe Partition Handling',
         '# With "pause_minority" mode:\n'
         '# - The minority partition (node1) will pause\n'
         '# - Majority partition (node2, node3) continues operating\n\n'
         '# With "autoheal" mode:\n'
         '# - Cluster detects partition\n'
         '# - Automatically restarts nodes to heal partition'),

        ('Step 5: Restore Network Connectivity',
         '# On node1, remove iptables rules:\n'
         'sudo iptables -D INPUT -s 192.168.1.102 -j DROP\n'
         'sudo iptables -D INPUT -s 192.168.1.103 -j DROP\n'
         'sudo iptables -D OUTPUT -d 192.168.1.102 -j DROP\n'
         'sudo iptables -D OUTPUT -d 192.168.1.103 -j DROP\n\n'
         '# Or flush all iptables rules:\n'
         'sudo iptables -F'),

        ('Step 6: Verify Partition Healing',
         '# Wait 30-60 seconds for autoheal to complete\n'
         'sudo rabbitmqctl cluster_status\n'
         '# All nodes should show as running again'),

        ('Step 7: Check Queue Integrity',
         'sudo rabbitmqctl list_queues name messages consumers\n'
         '# Verify queues are intact and synchronized'),

        ('Step 8: Verify Message Count',
         '# Compare message counts before and after partition\n'
         '# Check for any message loss or duplication')
    ]

    for step_title, step_cmd in partition_steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Expected Results:', bold=True)
    partition_results = [
        'Partition is detected within seconds on all nodes',
        'Partition handling mode activates automatically',
        'With autoheal: Nodes restart automatically to resolve partition',
        'With pause_minority: Minority partition pauses until network restored',
        'After healing, all nodes are running and synchronized',
        'No data corruption occurs',
        'Message loss may occur (messages in minority partition)',
        'Management UI shows partition events in logs'
    ]
    for result in partition_results:
        doc.add_paragraph(result, style='List Bullet')

    add_warning_box(doc,
        'Always use "autoheal" in production with odd number of nodes (3, 5, 7). '
        'Never use "ignore" mode as it can cause data inconsistency.')

    doc.add_page_break()

    # 9. Test Case 4: Graceful Node Shutdown
    add_heading(doc, '9. Test Case 4: Graceful Node Shutdown', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Verify that planned maintenance shutdown of a node allows proper queue migration '
        'and connection draining without message loss.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: LOW', bold=True)
    add_paragraph(doc, 'Expected Outcome: Zero message loss, graceful connection migration', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    graceful_steps = [
        ('Step 1: Prepare for Shutdown',
         '# Verify cluster health before shutdown\n'
         'sudo rabbitmqctl cluster_status\n'
         'sudo rabbitmqctl list_queues name messages'),

        ('Step 2: Stop New Connections',
         '# Block new connections to the node being shut down\n'
         'sudo rabbitmqctl suspend_listeners\n'
         '# Existing connections continue but no new ones accepted'),

        ('Step 3: Wait for Queues to Migrate',
         '# If node hosts queue masters, wait for clients to migrate\n'
         '# Monitor queue masters:\n'
         'watch -n 2 "sudo rabbitmqctl list_queues name pid"'),

        ('Step 4: Gracefully Stop RabbitMQ',
         '# Proper shutdown sequence\n'
         'sudo rabbitmqctl stop_app\n'
         'sudo rabbitmqctl stop'),

        ('Step 5: Verify Cluster Adapts',
         '# On remaining nodes:\n'
         'sudo rabbitmqctl cluster_status\n'
         '# Node should be marked as down gracefully'),

        ('Step 6: Verify Queue Availability',
         'sudo rabbitmqctl list_queues name messages consumers\n'
         '# All queues should be available on remaining nodes'),

        ('Step 7: Restart Node',
         'sudo systemctl start rabbitmq-server\n'
         '# Node should rejoin cluster cleanly'),

        ('Step 8: Resume Listeners',
         'sudo rabbitmqctl resume_listeners\n'
         '# Node can now accept new connections')
    ]

    for step_title, step_cmd in graceful_steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Expected Results:', bold=True)
    graceful_results = [
        'No message loss occurs',
        'Connections drain gracefully or reconnect to other nodes',
        'Queue masters migrate to other nodes if needed',
        'Cluster continues operating normally with remaining nodes',
        'Node rejoins cluster cleanly after restart',
        'No alarms or errors are triggered',
        'Complete operation within 2-3 minutes'
    ]
    for result in graceful_results:
        doc.add_paragraph(result, style='List Bullet')

    add_note_box(doc,
        'Graceful shutdown is the preferred method for maintenance. Always use this approach '
        'for planned maintenance activities.')

    doc.add_page_break()

    # 10. Test Case 5: Multiple Node Failure
    add_heading(doc, '10. Test Case 5: Multiple Node Failure', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Test cluster behavior when multiple nodes fail simultaneously or in quick succession. '
        'This is a catastrophic scenario requiring manual intervention.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: CRITICAL', bold=True)
    add_paragraph(doc, 'Expected Outcome: Cluster unavailable, manual recovery required', italic=True)

    doc.add_paragraph()
    add_warning_box(doc,
        'DANGER: This test will make the cluster completely unavailable. Only perform in '
        'test/staging environment. DO NOT run in production.')

    doc.add_paragraph()
    add_paragraph(doc, 'Scenario A: Two Nodes Fail (Quorum Lost)', bold=True)
    add_code_block(doc,
        '# Stop two nodes simultaneously\n'
        '# On node2:\n'
        'sudo systemctl stop rabbitmq-server\n\n'
        '# On node3:\n'
        'sudo systemctl stop rabbitmq-server\n\n'
        '# On node1 (remaining node):\n'
        'sudo rabbitmqctl cluster_status\n'
        '# Cluster loses quorum and becomes read-only or unavailable')

    doc.add_paragraph()
    add_paragraph(doc, 'Recovery Procedure:', bold=True)
    add_code_block(doc,
        '# Start nodes in sequence:\n'
        '# 1. Start the last node that was running (node1 in this case)\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# 2. Wait for node1 to start completely\n'
        'sudo rabbitmqctl wait /var/lib/rabbitmq/mnesia/rabbit@rabbitmq-node1.pid\n\n'
        '# 3. Start other nodes\n'
        '# On node2:\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# On node3:\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# 4. Verify all nodes are running\n'
        'sudo rabbitmqctl cluster_status')

    doc.add_paragraph()
    add_paragraph(doc, 'Scenario B: All Nodes Fail', bold=True)
    add_code_block(doc,
        '# If all nodes fail, perform these steps:\n\n'
        '# 1. Identify the last node to go down\n'
        '# Check system logs or timestamps\n\n'
        '# 2. Start the last node first\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# 3. If node fails to start, force boot:\n'
        'sudo rabbitmqctl force_boot\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# 4. Start remaining nodes in sequence\n\n'
        '# 5. Verify cluster status and queue integrity\n'
        'sudo rabbitmqctl cluster_status\n'
        'sudo rabbitmqctl list_queues name messages')

    add_warning_box(doc,
        'Using force_boot may result in data loss. Only use as last resort after consulting '
        'documentation and ensuring you understand the implications.')

    doc.add_page_break()

    # 11. Test Case 6: Disk Space Exhaustion
    add_heading(doc, '11. Test Case 6: Disk Space Exhaustion', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Verify RabbitMQ\'s behavior when disk space is exhausted and test the alarm system '
        'and recovery procedures.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: HIGH', bold=True)
    add_paragraph(doc, 'Expected Outcome: Disk alarm triggers, node blocks publishers', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    disk_steps = [
        ('Step 1: Check Current Disk Limits',
         'sudo rabbitmqctl eval "rabbit_disk_monitor:get_disk_free_limit()."\n'
         '# Or check in rabbitmq.conf:\n'
         '# disk_free_limit.absolute = 50GB'),

        ('Step 2: Simulate Disk Exhaustion',
         '# Create large file to fill disk space\n'
         'sudo fallocate -l 40G /tmp/large-file\n'
         '# Adjust size to trigger disk alarm based on your limit'),

        ('Step 3: Wait for Disk Alarm',
         '# Monitor alarms:\n'
         'watch -n 2 "sudo rabbitmqctl alarms"\n'
         '# Should show disk alarm: {resource_limit,disk,<node>}'),

        ('Step 4: Verify Publisher Blocking',
         '# Try to publish messages\n'
         '# Publishers should be blocked with error:\n'
         '# "ACCESS_REFUSED - disk alarm"'),

        ('Step 5: Check Management UI',
         '# Web UI should show red warning banner\n'
         '# Node status shows disk alarm active'),

        ('Step 6: Free Up Disk Space',
         'sudo rm /tmp/large-file\n'
         '# Or clean up logs:\n'
         'sudo journalctl --vacuum-time=7d'),

        ('Step 7: Verify Alarm Clears',
         'sudo rabbitmqctl alarms\n'
         '# Should return empty or no disk alarms'),

        ('Step 8: Resume Publishing',
         '# Publishers should automatically resume\n'
         '# Verify message flow restored')
    ]

    for step_title, step_cmd in disk_steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Expected Results:', bold=True)
    disk_results = [
        'Disk alarm triggers when free space drops below threshold',
        'Publishers are blocked from sending messages',
        'Consumers continue to drain queues',
        'Alarm clears automatically when disk space is freed',
        'Publishers resume automatically after alarm clears',
        'No message loss occurs (messages are rejected, not lost)',
        'Management UI displays alarm status clearly'
    ]
    for result in disk_results:
        doc.add_paragraph(result, style='List Bullet')

    doc.add_page_break()

    # 12. Test Case 7: Memory Pressure
    add_heading(doc, '12. Test Case 7: Memory Pressure', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Test RabbitMQ behavior under high memory usage and verify memory alarm and '
        'flow control mechanisms.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: HIGH', bold=True)
    add_paragraph(doc, 'Expected Outcome: Memory alarm triggers, publishers throttled', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    memory_steps = [
        ('Step 1: Check Memory Configuration',
         'sudo rabbitmqctl eval "vm_memory_monitor:get_memory_limit()."\n'
         '# Or check rabbitmq.conf:\n'
         '# vm_memory_high_watermark.relative = 0.4'),

        ('Step 2: Simulate High Memory Usage',
         '# Publish many large messages without consumers\n'
         'python3 producer.py --size 1MB --count 10000 --no-consumer'),

        ('Step 3: Monitor Memory Usage',
         'watch -n 2 "sudo rabbitmqctl status | grep memory"\n'
         '# Watch memory climb toward limit'),

        ('Step 4: Detect Memory Alarm',
         'sudo rabbitmqctl alarms\n'
         '# Should show: {resource_limit,memory,<node>}'),

        ('Step 5: Verify Publisher Throttling',
         '# Publishers should be blocked or slowed\n'
         '# Check for flow control in logs'),

        ('Step 6: Start Consumers to Drain Queues',
         'python3 consumer.py --queue test-queue --workers 10'),

        ('Step 7: Verify Memory Decreases',
         'watch -n 2 "sudo rabbitmqctl status | grep memory"\n'
         '# Memory should decrease as messages consumed'),

        ('Step 8: Verify Alarm Clears',
         'sudo rabbitmqctl alarms\n'
         '# Memory alarm should clear when below threshold')
    ]

    for step_title, step_cmd in memory_steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Expected Results:', bold=True)
    memory_results = [
        'Memory alarm triggers at configured watermark (default 40%)',
        'Publishers are throttled or blocked',
        'Consumers continue processing messages',
        'Memory usage decreases as queues drain',
        'Alarm clears when memory drops below watermark',
        'Publishers resume normal operation',
        'Node remains stable and responsive'
    ]
    for result in memory_results:
        doc.add_paragraph(result, style='List Bullet')

    add_note_box(doc,
        'Memory alarms are a sign of insufficient resources or unbalanced producer/consumer rates. '
        'Consider adding more consumers or increasing memory allocation.')

    doc.add_page_break()

    # 13. Test Case 8: Rolling Restart
    add_heading(doc, '13. Test Case 8: Rolling Restart', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Perform a rolling restart of all cluster nodes with zero downtime for upgrades '
        'or configuration changes.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: LOW', bold=True)
    add_paragraph(doc, 'Expected Outcome: Zero downtime, continuous message processing', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    rolling_steps = [
        ('Step 1: Verify Initial Cluster State',
         'sudo rabbitmqctl cluster_status\n'
         'sudo rabbitmqctl list_queues name messages consumers'),

        ('Step 2: Start Message Flow',
         '# Keep producers and consumers running throughout\n'
         'python3 producer.py --rate 100 &\n'
         'python3 consumer.py --workers 5 &'),

        ('Step 3: Restart Node 1 (Replica)',
         '# Choose a replica node first\n'
         'sudo systemctl restart rabbitmq-server\n'
         '# Wait for node to fully start\n'
         'sudo rabbitmqctl wait /var/lib/rabbitmq/mnesia/rabbit@rabbitmq-node1.pid'),

        ('Step 4: Verify Node 1 Rejoined',
         'sudo rabbitmqctl cluster_status\n'
         '# Verify all 3 nodes running before proceeding'),

        ('Step 5: Restart Node 2 (Replica)',
         '# On node2:\n'
         'sudo systemctl restart rabbitmq-server\n'
         'sudo rabbitmqctl wait /var/lib/rabbitmq/mnesia/rabbit@rabbitmq-node2.pid'),

        ('Step 6: Verify Node 2 Rejoined',
         'sudo rabbitmqctl cluster_status'),

        ('Step 7: Restart Node 3 (Last Node)',
         '# On node3:\n'
         'sudo systemctl restart rabbitmq-server\n'
         'sudo rabbitmqctl wait /var/lib/rabbitmq/mnesia/rabbit@rabbitmq-node3.pid'),

        ('Step 8: Verify Complete Cluster',
         'sudo rabbitmqctl cluster_status\n'
         '# All nodes should be running'),

        ('Step 9: Verify Message Flow',
         '# Check producer/consumer logs\n'
         '# Verify no message loss\n'
         '# Check for any connection interruptions')
    ]

    for step_title, step_cmd in rolling_steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Best Practices for Rolling Restart:', bold=True)
    rolling_practices = [
        'Always restart replica nodes before nodes hosting queue masters',
        'Wait for each node to fully rejoin before restarting the next',
        'Verify queue synchronization between restarts',
        'Monitor client connections and reconnection patterns',
        'Restart one node at a time (never multiple nodes simultaneously)',
        'Total time should be 5-10 minutes for 3-node cluster',
        'Schedule during low-traffic periods when possible'
    ]
    for practice in rolling_practices:
        doc.add_paragraph(practice, style='List Bullet')

    doc.add_page_break()

    # 14. Test Case 9: Queue Mirroring Validation
    add_heading(doc, '14. Test Case 9: Queue Mirroring Validation', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Verify that queue mirroring is working correctly and queues are synchronized '
        'across all nodes as configured.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: MEDIUM', bold=True)
    add_paragraph(doc, 'Expected Outcome: All queues properly mirrored and synchronized', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    mirror_steps = [
        ('Step 1: Check Mirroring Policy',
         'sudo rabbitmqctl list_policies\n'
         '# Verify ha-all or ha-exactly policy is active'),

        ('Step 2: Create Test Queues',
         'sudo rabbitmqadmin declare queue name=mirror-test-1 durable=true\n'
         'sudo rabbitmqadmin declare queue name=mirror-test-2 durable=true\n'
         'sudo rabbitmqadmin declare queue name=mirror-test-3 durable=true'),

        ('Step 3: Verify Queue Distribution',
         'sudo rabbitmqctl list_queues name pid slave_pids\n'
         '# Each queue should show 2 slave_pids (mirrors on other nodes)'),

        ('Step 4: Publish Messages to Test Queues',
         'for i in {1..1000}; do\n'
         '  sudo rabbitmqadmin publish routing_key=mirror-test-1 payload="test $i"\n'
         'done'),

        ('Step 5: Verify Message Counts on All Nodes',
         '# On each node, check message counts:\n'
         'sudo rabbitmqctl list_queues name messages\n'
         '# All nodes should report same counts'),

        ('Step 6: Check Synchronization Status',
         'sudo rabbitmqctl list_queues name synchronised_slave_pids\n'
         '# Should match slave_pids (all mirrors synchronized)'),

        ('Step 7: Test Failover',
         '# Identify and stop the master node for a queue\n'
         '# Verify messages are still available on mirrors'),

        ('Step 8: Verify Message Integrity',
         '# Consume all messages and verify count matches published\n'
         'python3 consumer.py --queue mirror-test-1 --count')
    ]

    for step_title, step_cmd in mirror_steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Mirroring Policies Explained:', bold=True)
    policies = [
        ('ha-all', 'Mirrors queue to all nodes in cluster (recommended for 3 nodes)'),
        ('ha-exactly:N', 'Mirrors queue to exactly N nodes (e.g., ha-exactly:2)'),
        ('ha-nodes', 'Mirrors queue to specific named nodes'),
        ('ha-sync-mode:automatic', 'New mirrors synchronize automatically'),
        ('ha-sync-mode:manual', 'New mirrors require manual sync trigger')
    ]
    for policy, description in policies:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(policy + ': ').bold = True
        para.add_run(description)

    doc.add_page_break()

    # 15. Test Case 10: Client Connection Failover
    add_heading(doc, '15. Test Case 10: Client Connection Failover', 1)

    add_paragraph(doc, 'Objective:', bold=True)
    add_paragraph(doc,
        'Verify that client applications can reconnect automatically when their connected '
        'node fails and continue processing messages.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Severity: MEDIUM', bold=True)
    add_paragraph(doc, 'Expected Outcome: Clients reconnect within 10 seconds', italic=True)

    doc.add_paragraph()
    add_paragraph(doc, 'Prerequisites:', bold=True)
    add_paragraph(doc,
        'Client application must be configured with multiple RabbitMQ node addresses for '
        'automatic failover.'
    )

    add_code_block(doc,
        '# Example client configuration (Python):\n'
        'connection_params = [\n'
        '    pika.ConnectionParameters(host="192.168.1.101"),\n'
        '    pika.ConnectionParameters(host="192.168.1.102"),\n'
        '    pika.ConnectionParameters(host="192.168.1.103")\n'
        ']')

    doc.add_paragraph()
    add_paragraph(doc, 'Test Steps:', bold=True)

    client_steps = [
        ('Step 1: Start Test Client',
         '# Run producer and consumer with multi-host config\n'
         'python3 producer.py --hosts 192.168.1.101,192.168.1.102,192.168.1.103'),

        ('Step 2: Verify Initial Connection',
         '# Check which node client is connected to\n'
         'sudo rabbitmqctl list_connections name peer_host peer_port state'),

        ('Step 3: Stop Connected Node',
         '# Identify which node client is using\n'
         '# Stop that specific node\n'
         'sudo systemctl stop rabbitmq-server'),

        ('Step 4: Monitor Client Reconnection',
         '# Watch client logs for reconnection attempts\n'
         '# Should see connection errors followed by successful reconnect'),

        ('Step 5: Verify Reconnection',
         '# On remaining nodes:\n'
         'sudo rabbitmqctl list_connections name peer_host\n'
         '# Client should appear connected to different node'),

        ('Step 6: Verify Message Flow Continues',
         '# Check producer/consumer continue operation\n'
         '# Verify message counts increase'),

        ('Step 7: Test Multiple Failovers',
         '# Stop second node\n'
         '# Verify client connects to third node\n'
         '# Test failover speed and reliability'),

        ('Step 8: Restore All Nodes',
         '# Start all nodes\n'
         '# Verify client maintains connection or rebalances')
    ]

    for step_title, step_cmd in client_steps:
        add_paragraph(doc, step_title, bold=True)
        add_code_block(doc, step_cmd)
        doc.add_paragraph()

    add_paragraph(doc, 'Expected Results:', bold=True)
    client_results = [
        'Client detects connection failure within 1-2 seconds',
        'Client attempts reconnection to alternative nodes',
        'Successful reconnection within 5-10 seconds',
        'Message processing resumes automatically',
        'No messages are lost (with proper acknowledgments)',
        'Client can handle multiple consecutive failovers',
        'Connection pools rebalance after nodes recover'
    ]
    for result in client_results:
        doc.add_paragraph(result, style='List Bullet')

    add_note_box(doc,
        'Clients must implement retry logic and configure multiple broker addresses. '
        'Use client library features for automatic failover (e.g., Pika\'s automatic '
        'connection recovery in Python).')

    doc.add_page_break()

    # 16. Monitoring and Validation Commands
    add_heading(doc, '16. Monitoring and Validation Commands', 1)

    add_paragraph(doc, 'Essential Monitoring Commands:', bold=True)

    commands = [
        ('Cluster Status', 'sudo rabbitmqctl cluster_status'),
        ('Node Health', 'sudo rabbitmq-diagnostics check_running\n'
                       'sudo rabbitmq-diagnostics check_local_alarms\n'
                       'sudo rabbitmq-diagnostics check_port_connectivity'),
        ('Queue Information', 'sudo rabbitmqctl list_queues name messages consumers memory slave_pids synchronised_slave_pids'),
        ('Connection Status', 'sudo rabbitmqctl list_connections name state user peer_host'),
        ('Channel Information', 'sudo rabbitmqctl list_channels connection number consumer_count'),
        ('Memory Usage', 'sudo rabbitmqctl status | grep -A 10 memory'),
        ('Disk Status', 'sudo rabbitmqctl status | grep -A 5 disk\n'
                       'df -h /var/lib/rabbitmq'),
        ('Active Alarms', 'sudo rabbitmqctl alarms'),
        ('Performance Metrics', 'sudo rabbitmqctl list_queues name messages_ready messages_unacknowledged message_stats'),
        ('Cluster-wide Stats', 'curl -u admin:password http://localhost:15672/api/overview')
    ]

    for cmd_name, cmd_text in commands:
        add_paragraph(doc, cmd_name + ':', bold=True)
        add_code_block(doc, cmd_text)
        doc.add_paragraph()

    add_paragraph(doc, 'Management UI Monitoring:', bold=True)
    ui_sections = [
        'Overview: Cluster-wide message rates, connections, channels',
        'Nodes: Individual node status, memory, disk, uptime',
        'Queues: Queue details, message rates, consumers, memory',
        'Connections: Active client connections and their channels',
        'Exchanges: Exchange types, bindings, message rates',
        'Admin > Policies: HA policies and their application'
    ]
    for section in ui_sections:
        doc.add_paragraph(section, style='List Bullet')

    doc.add_page_break()

    # 17. Common Issues and Troubleshooting
    add_heading(doc, '17. Common Issues and Troubleshooting', 1)

    issues = [
        ['Issue', 'Symptoms', 'Resolution'],
        ['Node won\'t join cluster', 'Error: inconsistent_cluster', 'Reset node: rabbitmqctl reset, then rejoin'],
        ['Queue not mirrored', 'slave_pids empty', 'Check HA policy, manually sync: rabbitmqctl sync_queue'],
        ['Split-brain not healing', 'Partition persists', 'Check partition_handling mode, manually restart nodes'],
        ['Memory alarm stuck', 'Alarm persists after cleanup', 'Restart node or set_vm_memory_high_watermark'],
        ['Slow failover', 'Takes > 60 seconds', 'Check network latency, reduce timeout settings'],
        ['Message loss on failover', 'Messages disappear', 'Implement publisher confirms and consumer acks'],
        ['Connections not rebalancing', 'All clients on one node', 'Restart clients or use load balancer'],
        ['Queues out of sync', 'synchronised_slave_pids empty', 'Trigger sync: rabbitmqctl sync_queue queue_name']
    ]
    add_table_with_header(doc, issues[0], issues[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Diagnostic Steps:', bold=True)
    diagnostic_steps = [
        'Check RabbitMQ logs: tail -f /var/log/rabbitmq/rabbit@hostname.log',
        'Check system logs: journalctl -u rabbitmq-server -f',
        'Verify network connectivity: ping and telnet to ports 5672, 25672, 4369',
        'Check Erlang cookie: cat /var/lib/rabbitmq/.erlang.cookie',
        'Verify hostname resolution: cat /etc/hosts',
        'Check resource usage: top, df -h, free -h',
        'Review cluster events in Management UI'
    ]
    for step in diagnostic_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_page_break()

    # 18. Recovery Procedures
    add_heading(doc, '18. Recovery Procedures', 1)

    add_paragraph(doc, 'Scenario: Node Won\'t Start After Crash', bold=True)
    add_code_block(doc,
        '# 1. Check logs for errors\n'
        'sudo journalctl -u rabbitmq-server -n 100\n\n'
        '# 2. Verify Erlang cookie permissions\n'
        'ls -la /var/lib/rabbitmq/.erlang.cookie\n'
        'sudo chmod 400 /var/lib/rabbitmq/.erlang.cookie\n'
        'sudo chown rabbitmq:rabbitmq /var/lib/rabbitmq/.erlang.cookie\n\n'
        '# 3. Try starting with force_boot if cluster start fails\n'
        'sudo rabbitmqctl force_boot\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# 4. If still fails, reset and rejoin cluster\n'
        'sudo rabbitmqctl stop_app\n'
        'sudo rabbitmqctl reset\n'
        'sudo rabbitmqctl join_cluster rabbit@rabbitmq-node1\n'
        'sudo rabbitmqctl start_app')

    doc.add_paragraph()
    add_paragraph(doc, 'Scenario: Persistent Network Partition', bold=True)
    add_code_block(doc,
        '# 1. Identify which nodes are in each partition\n'
        'sudo rabbitmqctl cluster_status\n\n'
        '# 2. Choose the partition with majority/most recent data\n\n'
        '# 3. Restart nodes in the minority partition\n'
        '# On minority nodes:\n'
        'sudo rabbitmqctl stop_app\n'
        'sudo rabbitmqctl start_app\n\n'
        '# 4. If partition persists, force healing:\n'
        '# On all nodes in minority partition:\n'
        'sudo rabbitmqctl stop_app\n'
        'sudo rabbitmqctl forget_cluster_node rabbit@minority-node\n'
        'sudo rabbitmqctl reset\n'
        'sudo rabbitmqctl join_cluster rabbit@majority-node\n'
        'sudo rabbitmqctl start_app')

    doc.add_paragraph()
    add_paragraph(doc, 'Scenario: All Nodes Down (Complete Cluster Failure)', bold=True)
    add_code_block(doc,
        '# 1. Identify the last node to go down (check logs/timestamps)\n\n'
        '# 2. Start that node first\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# 3. If it fails with "waiting for Mnesia tables":\n'
        'sudo rabbitmqctl force_boot\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# 4. Wait for first node to fully start\n'
        'sudo rabbitmqctl wait /var/lib/rabbitmq/mnesia/rabbit@node.pid\n\n'
        '# 5. Start other nodes one by one\n'
        'sudo systemctl start rabbitmq-server\n\n'
        '# 6. Verify cluster reformation\n'
        'sudo rabbitmqctl cluster_status\n\n'
        '# 7. Check queue integrity\n'
        'sudo rabbitmqctl list_queues name messages')

    doc.add_paragraph()
    add_paragraph(doc, 'Scenario: Queue Synchronization Issues', bold=True)
    add_code_block(doc,
        '# 1. Check which queues are out of sync\n'
        'sudo rabbitmqctl list_queues name slave_pids synchronised_slave_pids\n\n'
        '# 2. Manually trigger synchronization\n'
        'sudo rabbitmqctl sync_queue queue_name\n\n'
        '# 3. Or sync all queues:\n'
        'for queue in $(sudo rabbitmqctl list_queues -q name); do\n'
        '    sudo rabbitmqctl sync_queue "$queue"\n'
        'done\n\n'
        '# 4. Monitor sync progress\n'
        'watch -n 2 "sudo rabbitmqctl list_queues name synchronised_slave_pids"')

    doc.add_page_break()

    # 19. Best Practices for Production
    add_heading(doc, '19. Best Practices for Production', 1)

    practices = [
        ('Cluster Configuration', [
            'Use odd number of nodes (3, 5, 7) for proper quorum',
            'Set cluster_partition_handling to "autoheal"',
            'Configure HA policy for critical queues',
            'Enable automatic queue synchronization',
            'Use durable queues for persistent messaging'
        ]),
        ('Monitoring', [
            'Implement comprehensive monitoring (Prometheus + Grafana)',
            'Set up alerts for disk/memory alarms',
            'Monitor queue depth and consumer lag',
            'Track connection and channel counts',
            'Log cluster events and partition occurrences',
            'Monitor message rates and throughput'
        ]),
        ('Client Applications', [
            'Configure clients with all cluster node addresses',
            'Implement automatic reconnection logic',
            'Use publisher confirms for critical messages',
            'Always acknowledge messages after processing',
            'Set appropriate prefetch counts',
            'Handle connection failures gracefully'
        ]),
        ('Resource Management', [
            'Set appropriate memory and disk limits',
            'Configure memory_high_watermark (40% default)',
            'Set disk_free_limit to adequate value',
            'Monitor and rotate logs regularly',
            'Plan for 2x peak load capacity',
            'Use SSD for persistence files'
        ]),
        ('Disaster Recovery', [
            'Regular backup of queue definitions',
            'Document recovery procedures',
            'Test failover scenarios quarterly',
            'Maintain separate staging environment for testing',
            'Keep configuration in version control',
            'Document node startup order'
        ]),
        ('Maintenance', [
            'Use rolling restarts for updates',
            'Schedule maintenance during low-traffic periods',
            'Test configuration changes in staging first',
            'Keep RabbitMQ and Erlang versions up to date',
            'Plan upgrade paths carefully',
            'Notify clients before maintenance windows'
        ])
    ]

    for category, items in practices:
        add_paragraph(doc, category + ':', bold=True)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # 20. Test Results Template
    add_heading(doc, '20. Test Results Template', 1)

    add_paragraph(doc, 'Test Execution Record:', bold=True)
    test_info = [
        ['Field', 'Value'],
        ['Test Date', ''],
        ['Tester Name', ''],
        ['Environment', 'Test / Staging / Production'],
        ['RabbitMQ Version', ''],
        ['Cluster Configuration', '3 nodes / Other'],
        ['HA Policy', 'ha-all / ha-exactly:2']
    ]
    add_table_with_header(doc, test_info[0], test_info[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Test Case Results:', bold=True)
    results_template = [
        ['Test Case', 'Status', 'Duration', 'Notes'],
        ['TC1: Master Node Failure', 'Pass/Fail', '', ''],
        ['TC2: Replica Node Failure', 'Pass/Fail', '', ''],
        ['TC3: Network Partition', 'Pass/Fail', '', ''],
        ['TC4: Graceful Shutdown', 'Pass/Fail', '', ''],
        ['TC5: Multiple Node Failure', 'Pass/Fail', '', ''],
        ['TC6: Disk Space Exhaustion', 'Pass/Fail', '', ''],
        ['TC7: Memory Pressure', 'Pass/Fail', '', ''],
        ['TC8: Rolling Restart', 'Pass/Fail', '', ''],
        ['TC9: Queue Mirroring', 'Pass/Fail', '', ''],
        ['TC10: Client Failover', 'Pass/Fail', '', '']
    ]
    add_table_with_header(doc, results_template[0], results_template[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Success Criteria:', bold=True)
    criteria = [
        'All automated failovers complete within expected time',
        'Message loss is zero or within acceptable limits (<0.1%)',
        'Cluster recovers to full health after each test',
        'No manual intervention required for automatic scenarios',
        'Clients reconnect successfully within 10 seconds',
        'Monitoring and alerting triggered appropriately',
        'All recovery procedures documented and tested'
    ]
    for criterion in criteria:
        doc.add_paragraph(criterion, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Issues Found:', bold=True)
    issues_template = [
        ['Issue #', 'Description', 'Severity', 'Resolution'],
        ['1', '', 'High/Medium/Low', ''],
        ['2', '', 'High/Medium/Low', ''],
        ['3', '', 'High/Medium/Low', '']
    ]
    add_table_with_header(doc, issues_template[0], issues_template[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Overall Assessment:', bold=True)
    add_paragraph(doc, '[Pass / Fail / Pass with Conditions]')
    doc.add_paragraph()
    add_paragraph(doc, 'Recommendations:', bold=True)
    add_paragraph(doc, '[List any recommendations for improvement]')

    doc.add_page_break()

    # Appendix
    add_heading(doc, 'Appendix: Sample Test Scripts', 1)

    add_paragraph(doc, 'Sample Producer Script (producer.py):', bold=True)
    add_code_block(doc,
        '#!/usr/bin/env python3\n'
        'import pika\n'
        'import time\n'
        'import sys\n\n'
        'def main():\n'
        '    # Multi-host failover configuration\n'
        '    hosts = ["192.168.1.101", "192.168.1.102", "192.168.1.103"]\n'
        '    credentials = pika.PlainCredentials("admin", "password")\n'
        '    \n'
        '    params = [\n'
        '        pika.ConnectionParameters(\n'
        '            host=h, credentials=credentials,\n'
        '            heartbeat=30, blocked_connection_timeout=300\n'
        '        ) for h in hosts\n'
        '    ]\n'
        '    \n'
        '    connection = pika.BlockingConnection(params)\n'
        '    channel = connection.channel()\n'
        '    \n'
        '    # Enable publisher confirms\n'
        '    channel.confirm_delivery()\n'
        '    \n'
        '    message_count = 0\n'
        '    try:\n'
        '        while True:\n'
        '            message = f"Message {message_count}"\n'
        '            channel.basic_publish(\n'
        '                exchange="",\n'
        '                routing_key="test-failover-queue",\n'
        '                body=message,\n'
        '                properties=pika.BasicProperties(delivery_mode=2)\n'
        '            )\n'
        '            print(f"Sent: {message}")\n'
        '            message_count += 1\n'
        '            time.sleep(0.1)\n'
        '    except KeyboardInterrupt:\n'
        '        connection.close()\n\n'
        'if __name__ == "__main__":\n'
        '    main()')

    doc.add_paragraph()
    add_paragraph(doc, 'Sample Consumer Script (consumer.py):', bold=True)
    add_code_block(doc,
        '#!/usr/bin/env python3\n'
        'import pika\n'
        'import time\n\n'
        'def callback(ch, method, properties, body):\n'
        '    print(f"Received: {body.decode()}")\n'
        '    # Simulate processing\n'
        '    time.sleep(0.05)\n'
        '    # Acknowledge message\n'
        '    ch.basic_ack(delivery_tag=method.delivery_tag)\n\n'
        'def main():\n'
        '    hosts = ["192.168.1.101", "192.168.1.102", "192.168.1.103"]\n'
        '    credentials = pika.PlainCredentials("admin", "password")\n'
        '    \n'
        '    params = [\n'
        '        pika.ConnectionParameters(\n'
        '            host=h, credentials=credentials, heartbeat=30\n'
        '        ) for h in hosts\n'
        '    ]\n'
        '    \n'
        '    connection = pika.BlockingConnection(params)\n'
        '    channel = connection.channel()\n'
        '    \n'
        '    # Set prefetch count\n'
        '    channel.basic_qos(prefetch_count=10)\n'
        '    \n'
        '    channel.basic_consume(\n'
        '        queue="test-failover-queue",\n'
        '        on_message_callback=callback\n'
        '    )\n'
        '    \n'
        '    print("Waiting for messages...")\n'
        '    try:\n'
        '        channel.start_consuming()\n'
        '    except KeyboardInterrupt:\n'
        '        channel.stop_consuming()\n'
        '        connection.close()\n\n'
        'if __name__ == "__main__":\n'
        '    main()')

    doc.add_page_break()

    # Footer
    add_heading(doc, 'Document Information', 1)
    doc_info = [
        ['Field', 'Value'],
        ['Document Version', '1.0'],
        ['Last Updated', '2026-01-14'],
        ['Document Type', 'Test Procedures and Scenarios'],
        ['Target Audience', 'DevOps Engineers, SREs, System Administrators'],
        ['Classification', 'Internal Technical Documentation']
    ]
    add_table_with_header(doc, doc_info[0], doc_info[1:])

    return doc

if __name__ == '__main__':
    doc = create_failover_document()
    doc.save('RabbitMQ_Cluster_Failover_Test_Cases.docx')
    print("RabbitMQ failover scenarios documentation created successfully!")
