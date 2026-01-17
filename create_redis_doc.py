#!/usr/bin/env python3
"""
Generate Redis 8.x Documentation for RHEL 8
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def add_code_block(doc, text):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.style = 'Intense Quote'
    return para

def add_table_with_header(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    return table

def create_redis_document():
    """Create the Redis documentation"""
    doc = Document()

    # Title Page
    title = doc.add_heading('Redis 8.x', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Installation and Configuration Guide', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    platform = doc.add_heading('Red Hat Enterprise Linux 8', 3)
    platform.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Purpose and Use Cases',
        '3. System Prerequisites',
        '4. Operating System Requirements',
        '5. Standalone Redis Installation',
        '6. Redis Sentinel Configuration (High Availability)',
        '7. Redis Cluster Configuration',
        '8. Post-Installation Configuration',
        '9. Security Hardening',
        '10. Verification and Testing',
        '11. Performance Tuning',
        '12. Best Practices',
        '13. Troubleshooting'
    ]
    for item in toc_items:
        add_paragraph(doc, item)

    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc,
        'Redis (Remote Dictionary Server) is an open-source, in-memory data structure store '
        'that can be used as a database, cache, message broker, and streaming engine. Redis 8.x '
        'introduces significant performance improvements, enhanced clustering capabilities, and '
        'new data structures that make it more powerful and efficient than ever before.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Key Features:', bold=True)
    features = [
        'In-memory data storage with optional persistence',
        'Support for multiple data structures (strings, hashes, lists, sets, sorted sets, bitmaps, hyperloglogs, streams)',
        'Built-in replication and high availability',
        'Automatic partitioning with Redis Cluster',
        'Pub/Sub messaging',
        'Lua scripting support',
        'Transactions and pipeline support',
        'JSON document storage with RedisJSON module'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # 2. Purpose and Use Cases
    add_heading(doc, '2. Purpose and Use Cases', 1)

    add_paragraph(doc, 'Redis Purpose:', bold=True)
    add_paragraph(doc,
        'Redis serves as a high-performance data store that provides sub-millisecond response '
        'times, enabling millions of requests per second. Its versatile data structures and '
        'atomic operations make it ideal for real-time applications requiring fast data access.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Common Use Cases:', bold=True)

    use_cases = [
        ('Caching', 'Store frequently accessed data to reduce database load and improve application response times'),
        ('Session Management', 'Store user session data with automatic expiration for web applications'),
        ('Real-time Analytics', 'Process and analyze streaming data with low latency using Redis Streams'),
        ('Message Queues', 'Implement job queues and background processing with lists and sorted sets'),
        ('Leaderboards and Counters', 'Maintain real-time rankings and statistics using sorted sets'),
        ('Rate Limiting', 'Control API request rates using counters with expiration'),
        ('Geospatial Applications', 'Store and query location-based data with geospatial indexes'),
        ('Pub/Sub Messaging', 'Build real-time notification and chat systems'),
        ('Full-text Search', 'Implement search functionality with RediSearch module'),
        ('Time Series Data', 'Store and query time-series data efficiently with RedisTimeSeries')
    ]

    for title, description in use_cases:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(title + ': ').bold = True
        para.add_run(description)

    doc.add_page_break()

    # 3. System Prerequisites
    add_heading(doc, '3. System Prerequisites', 1)

    add_paragraph(doc, 'Hardware Requirements:', bold=True)
    hw_reqs = [
        ['Component', 'Minimum', 'Recommended'],
        ['CPU', '2 Cores', '4+ Cores'],
        ['RAM', '4 GB', '16 GB+ (depends on dataset size)'],
        ['Disk Space', '10 GB', '50 GB+ (for persistence and logs)'],
        ['Network', '1 Gbps', '10 Gbps (for high throughput)']
    ]
    add_table_with_header(doc, hw_reqs[0], hw_reqs[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Important Notes:', bold=True)
    notes = [
        'Redis is primarily memory-based; ensure sufficient RAM for your dataset',
        'Use SSDs for persistence (RDB/AOF) to improve I/O performance',
        'For clusters, use dedicated network interfaces for inter-node communication'
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Software Requirements:', bold=True)
    sw_items = [
        'Red Hat Enterprise Linux 8.x (8.5 or later recommended)',
        'GCC compiler and build tools',
        'Root or sudo access for installation',
        'SELinux configured appropriately',
        'Firewall rules configured for Redis ports'
    ]
    for item in sw_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Network Requirements:', bold=True)
    network_items = [
        'Port 6379: Default Redis server port',
        'Port 16379: Redis Cluster bus port (client port + 10000)',
        'Port 26379: Redis Sentinel default port',
        'Additional ports as configured for multiple instances'
    ]
    for item in network_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 4. Operating System Requirements
    add_heading(doc, '4. Operating System Requirements and Configuration', 1)

    add_paragraph(doc,
        'This section provides comprehensive operating system configuration requirements for '
        'deploying Redis 8.x in production environments on RHEL 8. These configurations have been '
        'validated for enterprise deployments and follow industry best practices.'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.1 System Updates and Patch Management', bold=True)
    add_paragraph(doc,
        'Purpose: Maintain system security and stability through regular patching. Redis performance '
        'benefits from recent kernel improvements. Coordinate updates with change management processes '
        'and schedule during planned maintenance windows.'
    )
    add_code_block(doc,
        '# Verify current OS version\n'
        'cat /etc/redhat-release\n\n'
        '# Check for available updates\n'
        'sudo dnf check-update\n\n'
        '# Apply system updates (schedule during maintenance window)\n'
        'sudo dnf update -y\n\n'
        '# Check if reboot is required after kernel updates\n'
        'sudo needs-restarting -r\n\n'
        '# If reboot is required:\n'
        'sudo reboot'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.2 Install Development Tools and Dependencies', bold=True)
    add_paragraph(doc,
        'Purpose: Redis 8.x is typically compiled from source to ensure the latest features and '
        'optimizations. Install GCC, make, and other build tools required for compilation. These '
        'tools are also useful for building Redis modules and troubleshooting.'
    )
    add_code_block(doc,
        '# Install development tools group\n'
        'sudo dnf groupinstall -y "Development Tools"\n\n'
        '# Install specific build dependencies\n'
        'sudo dnf install -y gcc gcc-c++ make wget curl tcl jemalloc-devel\n\n'
        '# Install troubleshooting and monitoring tools\n'
        'sudo dnf install -y net-tools telnet nc bind-utils\n'
        'sudo dnf install -y sysstat htop iotop\n\n'
        '# Verify GCC installation\n'
        'gcc --version\n\n'
        '# Verify make installation\n'
        'make --version'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.3 Disable Transparent Huge Pages (THP)', bold=True)
    add_paragraph(doc,
        'Purpose: Transparent Huge Pages can cause significant latency spikes in Redis operations. '
        'The Linux kernel\'s THP feature attempts to optimize memory management but can interfere '
        'with Redis\' memory access patterns, causing delays during background saves (BGSAVE/BGREWRITE). '
        'Disabling THP is MANDATORY for production Redis deployments and is recommended by Redis Labs.'
    )
    add_code_block(doc,
        '# Check current THP status\n'
        'cat /sys/kernel/mm/transparent_hugepage/enabled\n'
        '# Current setting will be in [brackets]\n\n'
        '# Disable THP immediately (takes effect right away)\n'
        'echo never | sudo tee /sys/kernel/mm/transparent_hugepage/enabled\n'
        'echo never | sudo tee /sys/kernel/mm/transparent_hugepage/defrag\n\n'
        '# Method 1: Make persistent using rc.local\n'
        'sudo vi /etc/rc.d/rc.local\n'
        '# Add these lines:\n'
        '#!/bin/bash\n'
        'echo never > /sys/kernel/mm/transparent_hugepage/enabled\n'
        'echo never > /sys/kernel/mm/transparent_hugepage/defrag\n\n'
        '# Make rc.local executable\n'
        'sudo chmod +x /etc/rc.d/rc.local\n'
        'sudo systemctl enable rc-local\n\n'
        '# Method 2: Use systemd service (recommended)\n'
        'sudo vi /etc/systemd/system/disable-thp.service\n'
        '# Add:\n'
        '[Unit]\n'
        'Description=Disable Transparent Huge Pages (THP)\n'
        'DefaultDependencies=no\n'
        'After=sysinit.target local-fs.target\n'
        'Before=redis.service\n\n'
        '[Service]\n'
        'Type=oneshot\n'
        'ExecStart=/bin/sh -c "echo never > /sys/kernel/mm/transparent_hugepage/enabled"\n'
        'ExecStart=/bin/sh -c "echo never > /sys/kernel/mm/transparent_hugepage/defrag"\n\n'
        '[Install]\n'
        'WantedBy=basic.target\n\n'
        '# Enable the service\n'
        'sudo systemctl daemon-reload\n'
        'sudo systemctl enable disable-thp.service\n'
        'sudo systemctl start disable-thp.service\n\n'
        '# Verify THP is disabled\n'
        'cat /sys/kernel/mm/transparent_hugepage/enabled\n'
        '# Should show: always madvise [never]'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.4 Configure Kernel Parameters', bold=True)
    add_paragraph(doc,
        'Purpose: Optimize Linux kernel parameters for Redis workloads. These settings improve '
        'Redis performance, stability, and reliability by adjusting memory management, network '
        'stack behavior, and file handling. Changes are persistent across reboots.'
    )
    add_code_block(doc,
        '# Create Redis-specific sysctl configuration\n'
        'sudo vi /etc/sysctl.d/99-redis.conf\n'
        '# Add the following parameters:\n\n'
        '# Memory overcommit for Redis background saves\n'
        '# Required for BGSAVE and BGREWRITE operations\n'
        'vm.overcommit_memory = 1\n\n'
        '# Increase connection backlog queue\n'
        '# Prevents "connection refused" errors under high load\n'
        'net.core.somaxconn = 65535\n\n'
        '# Increase maximum file descriptors\n'
        'fs.file-max = 100000\n\n'
        '# Disable swapping for better performance\n'
        '# WARNING: Ensure sufficient RAM before setting\n'
        'vm.swappiness = 0\n\n'
        '# TCP settings for better network performance\n'
        'net.ipv4.tcp_max_syn_backlog = 4096\n'
        'net.ipv4.tcp_keepalive_time = 300\n'
        'net.ipv4.tcp_keepalive_intvl = 30\n'
        'net.ipv4.tcp_keepalive_probes = 3\n\n'
        '# Apply all changes immediately\n'
        'sudo sysctl -p /etc/sysctl.d/99-redis.conf\n\n'
        '# Verify settings\n'
        'sudo sysctl vm.overcommit_memory\n'
        'sudo sysctl net.core.somaxconn\n'
        'sudo sysctl fs.file-max'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Parameter Explanations:', bold=True)
    params_detailed = [
        ('vm.overcommit_memory = 1', 'Allows Redis to fork child processes for background saves '
         'without failing. Essential for RDB snapshots and AOF rewrites.'),
        ('net.core.somaxconn = 65535', 'Increases connection queue size to handle burst traffic. '
         'Default (128) is too low for high-traffic Redis instances.'),
        ('fs.file-max = 100000', 'Sets system-wide maximum file descriptor limit. Supports high '
         'connection counts and multiple Redis instances.'),
        ('vm.swappiness = 0', 'Prevents Redis memory from being swapped to disk, which causes severe '
         'latency. Only set if sufficient RAM is available.'),
        ('net.ipv4.tcp_keepalive_time = 300', 'Detects dead connections faster, freeing resources. '
         'Reduces idle connection overhead.')
    ]
    for param_name, param_desc in params_detailed:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(param_name + ': ').bold = True
        para.add_run(param_desc)

    doc.add_paragraph()
    add_paragraph(doc, '4.5 Configure System Limits', bold=True)
    add_paragraph(doc,
        'Purpose: Redis can maintain thousands of client connections simultaneously. Default Linux '
        'limits (1024 open files) are insufficient for production deployments. Increase file '
        'descriptor limits to prevent "Too many open files" errors.'
    )
    add_code_block(doc,
        '# Configure user limits in limits.conf\n'
        'sudo vi /etc/security/limits.conf\n'
        '# Add at the end of file:\n'
        'redis soft nofile 65536\n'
        'redis hard nofile 65536\n'
        'redis soft nproc 4096\n'
        'redis hard nproc 4096\n\n'
        '# Configure systemd service limits\n'
        'sudo mkdir -p /etc/systemd/system/redis.service.d/\n'
        'sudo vi /etc/systemd/system/redis.service.d/limits.conf\n'
        '# Add:\n'
        '[Service]\n'
        'LimitNOFILE=65536\n'
        'LimitNPROC=4096\n\n'
        '# Reload systemd to apply changes\n'
        'sudo systemctl daemon-reload\n\n'
        '# After Redis starts, verify limits:\n'
        '# cat /proc/$(pidof redis-server)/limits | grep "open files"\n'
        '# Should show: Max open files 65536'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.6 Firewall Configuration', bold=True)
    add_paragraph(doc,
        'Purpose: Configure firewalld to allow Redis traffic while maintaining security posture. '
        'In production environments, restrict Redis access to application servers only. Never expose '
        'Redis directly to the internet. Use firewall zones and rich rules for granular access control.'
    )
    add_code_block(doc,
        '# Check firewall status\n'
        'sudo firewall-cmd --state\n'
        'sudo firewall-cmd --get-active-zones\n\n'
        '# Add Redis ports to firewall (permanent configuration)\n'
        'sudo firewall-cmd --permanent --add-port=6379/tcp   # Redis server\n'
        'sudo firewall-cmd --permanent --add-port=16379/tcp  # Cluster bus\n'
        'sudo firewall-cmd --permanent --add-port=26379/tcp  # Sentinel\n\n'
        '# PRODUCTION: Restrict access to specific application servers\n'
        '# Example - Allow only from app server subnet 10.10.20.0/24:\n'
        'sudo firewall-cmd --permanent --add-rich-rule=\'rule family="ipv4" \\\n'
        '  source address="10.10.20.0/24" port port="6379" protocol="tcp" accept\'\n\n'
        '# For multiple application servers, add separate rules:\n'
        '# sudo firewall-cmd --permanent --add-rich-rule=\'rule family="ipv4" \\\n'
        '#   source address="10.10.30.0/24" port port="6379" protocol="tcp" accept\'\n\n'
        '# Remove the generic port rule if using rich rules:\n'
        '# sudo firewall-cmd --permanent --remove-port=6379/tcp\n\n'
        '# Apply firewall changes\n'
        'sudo firewall-cmd --reload\n\n'
        '# List all configured rules\n'
        'sudo firewall-cmd --list-all\n\n'
        '# Test connectivity from application server\n'
        'telnet redis-server 6379\n'
        'nc -zv redis-server 6379'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.7 SELinux Configuration', bold=True)
    add_paragraph(doc,
        'Purpose: Configure SELinux to allow Redis operations while maintaining mandatory access '
        'controls. For production environments, keep SELinux in enforcing mode and configure appropriate '
        'policies. Disabling SELinux violates security best practices and may fail compliance audits '
        '(PCI-DSS, HIPAA, SOC2). Work with your security team to create custom policies if needed.'
    )
    add_code_block(doc,
        '# Check current SELinux status\n'
        'sestatus\n'
        'getenforce\n\n'
        '# PRODUCTION RECOMMENDED: Configure SELinux policies\n'
        '# Install SELinux policy tools\n'
        'sudo dnf install -y policycoreutils-python-utils\n\n'
        '# Allow Redis to bind to network ports\n'
        'sudo semanage port -a -t redis_port_t -p tcp 6379\n'
        'sudo semanage port -a -t redis_port_t -p tcp 16379\n'
        'sudo semanage port -a -t redis_port_t -p tcp 26379\n\n'
        '# Set correct SELinux context for Redis directories\n'
        'sudo semanage fcontext -a -t redis_var_lib_t "/var/lib/redis(/.*)?"\n'
        'sudo semanage fcontext -a -t redis_log_t "/var/log/redis(/.*)?"\n'
        'sudo restorecon -Rv /var/lib/redis\n'
        'sudo restorecon -Rv /var/log/redis\n\n'
        '# Enable required boolean for network access\n'
        'sudo setsebool -P redis_enable_homedirs 1\n\n'
        '# Monitor for SELinux denials\n'
        'sudo ausearch -m avc -ts recent | grep redis\n\n'
        '# If denials occur, generate custom policy:\n'
        '# 1. Reproduce the denied operation\n'
        '# 2. Generate policy module:\n'
        '#    sudo grep redis /var/log/audit/audit.log | audit2allow -M redis_custom\n'
        '# 3. Review the generated policy:\n'
        '#    cat redis_custom.te\n'
        '# 4. Install policy (after security team approval):\n'
        '#    sudo semodule -i redis_custom.pp\n\n'
        '# DEVELOPMENT/TESTING ONLY: Temporary permissive mode\n'
        '# sudo setenforce 0  # Temporary until reboot\n\n'
        '# NEVER DO THIS IN PRODUCTION: Permanently disable\n'
        '# sudo sed -i "s/SELINUX=enforcing/SELINUX=disabled/" /etc/selinux/config\n'
        '# Requires reboot and violates security policies'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.8 Time Synchronization (NTP/Chrony)', bold=True)
    add_paragraph(doc,
        'Purpose: Accurate time synchronization is critical for Redis Cluster, Sentinel, and logging. '
        'Time drift between nodes can cause split-brain scenarios, certificate validation failures, '
        'and difficulty in log correlation. Synchronize all Redis nodes with corporate NTP servers.'
    )
    add_code_block(doc,
        '# Install chrony (modern NTP client)\n'
        'sudo dnf install -y chrony\n\n'
        '# Configure NTP servers\n'
        'sudo vi /etc/chrony.conf\n'
        '# Replace default servers with corporate NTP servers:\n'
        '# server ntp1.corp.example.com iburst\n'
        '# server ntp2.corp.example.com iburst\n'
        '# server ntp3.corp.example.com iburst\n\n'
        '# Enable and start chrony\n'
        'sudo systemctl enable chronyd\n'
        'sudo systemctl restart chronyd\n\n'
        '# Verify synchronization status\n'
        'chronyc tracking\n'
        'chronyc sources -v\n\n'
        '# Check if system clock is synchronized\n'
        'timedatectl\n\n'
        '# View time on all Redis nodes to verify sync\n'
        'date\n\n'
        '# Maximum time drift should be < 100ms'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.9 Hostname Configuration (For Cluster/Sentinel)', bold=True)
    add_paragraph(doc,
        'Purpose: Redis Sentinel and Cluster require proper hostname resolution. Configure FQDNs '
        'and ensure all nodes can resolve each other. Use corporate DNS in production; /etc/hosts '
        'is acceptable for testing but not recommended for production due to management overhead.'
    )
    add_code_block(doc,
        '# Set FQDN hostname\n'
        'sudo hostnamectl set-hostname redis-node1.example.com\n\n'
        '# Verify hostname\n'
        'hostname\n'
        'hostname -f\n\n'
        '# Configure /etc/hosts if DNS is unavailable\n'
        'sudo vi /etc/hosts\n'
        '# Add all Redis nodes:\n'
        '10.10.10.101 redis-node1.example.com redis-node1\n'
        '10.10.10.102 redis-node2.example.com redis-node2\n'
        '10.10.10.103 redis-node3.example.com redis-node3\n\n'
        '# Test resolution from each node\n'
        'ping -c 2 redis-node1.example.com\n'
        'ping -c 2 redis-node2.example.com\n'
        'ping -c 2 redis-node3.example.com\n\n'
        '# Verify reverse lookup\n'
        'nslookup redis-node1.example.com'
    )

    doc.add_page_break()

    # 5. Standalone Redis Installation
    add_heading(doc, '5. Standalone Redis Installation', 1)

    add_paragraph(doc, 'Method 1: Install from Source (Recommended for Redis 8.x)', bold=True)
    doc.add_paragraph()

    add_paragraph(doc, 'Step 1: Download Redis Source', bold=True)
    add_code_block(doc,
        'cd /tmp\n'
        'wget https://download.redis.io/redis-stable.tar.gz\n'
        '# Or download specific version:\n'
        'wget https://github.com/redis/redis/archive/8.0.0.tar.gz\n'
        'tar -xzvf redis-stable.tar.gz\n'
        'cd redis-stable'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 2: Compile Redis', bold=True)
    add_paragraph(doc, 'Purpose: Build Redis binaries from source.')
    add_code_block(doc,
        'make\n'
        'make test  # Optional but recommended\n'
        'sudo make install'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Create Redis User and Directories', bold=True)
    add_code_block(doc,
        '# Create redis user\n'
        'sudo useradd -r -s /bin/false redis\n\n'
        '# Create directories\n'
        'sudo mkdir -p /etc/redis\n'
        'sudo mkdir -p /var/lib/redis\n'
        'sudo mkdir -p /var/log/redis\n'
        'sudo mkdir -p /var/run/redis\n\n'
        '# Set ownership\n'
        'sudo chown -R redis:redis /var/lib/redis\n'
        'sudo chown -R redis:redis /var/log/redis\n'
        'sudo chown -R redis:redis /var/run/redis'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 4: Configure Redis', bold=True)
    add_code_block(doc,
        '# Copy default configuration\n'
        'sudo cp redis.conf /etc/redis/redis.conf\n\n'
        '# Edit configuration\n'
        'sudo vi /etc/redis/redis.conf'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Essential Configuration Parameters:', bold=True)
    add_code_block(doc,
        'bind 0.0.0.0\n'
        'protected-mode yes\n'
        'port 6379\n'
        'tcp-backlog 511\n'
        'timeout 0\n'
        'tcp-keepalive 300\n'
        'daemonize yes\n'
        'supervised systemd\n'
        'pidfile /var/run/redis/redis-server.pid\n'
        'loglevel notice\n'
        'logfile /var/log/redis/redis-server.log\n'
        'databases 16\n'
        'save 900 1\n'
        'save 300 10\n'
        'save 60 10000\n'
        'dir /var/lib/redis\n'
        'requirepass YourSecurePasswordHere\n'
        'maxmemory 2gb\n'
        'maxmemory-policy allkeys-lru'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 5: Create Systemd Service', bold=True)
    add_paragraph(doc, 'Purpose: Manage Redis as a systemd service.')
    add_code_block(doc,
        'sudo vi /etc/systemd/system/redis.service\n\n'
        '# Add the following content:\n'
        '[Unit]\n'
        'Description=Redis In-Memory Data Store\n'
        'After=network.target\n\n'
        '[Service]\n'
        'Type=forking\n'
        'User=redis\n'
        'Group=redis\n'
        'ExecStart=/usr/local/bin/redis-server /etc/redis/redis.conf\n'
        'ExecStop=/usr/local/bin/redis-cli shutdown\n'
        'Restart=always\n'
        'RestartSec=5\n'
        'LimitNOFILE=65536\n\n'
        '[Install]\n'
        'WantedBy=multi-user.target'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 6: Start Redis Service', bold=True)
    add_code_block(doc,
        'sudo systemctl daemon-reload\n'
        'sudo systemctl enable redis\n'
        'sudo systemctl start redis\n'
        'sudo systemctl status redis'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Method 2: Install from EPEL Repository', bold=True)
    add_code_block(doc,
        '# Enable EPEL repository\n'
        'sudo dnf install -y epel-release\n\n'
        '# Install Redis\n'
        'sudo dnf install -y redis\n\n'
        '# Note: EPEL may not have Redis 8.x yet\n'
        '# Verify version:\n'
        'redis-server --version'
    )

    doc.add_page_break()

    # 6. Redis Sentinel Configuration
    add_heading(doc, '6. Redis Sentinel Configuration (High Availability)', 1)

    add_paragraph(doc,
        'Redis Sentinel provides high availability for Redis by monitoring master and replica '
        'instances, performing automatic failover when the master fails, and serving as a '
        'configuration provider for clients.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Architecture Overview:', bold=True)
    arch_points = [
        'At least 3 Sentinel instances recommended (for quorum)',
        'One Redis master and multiple replicas',
        'Sentinels monitor master and replicas',
        'Automatic failover when master becomes unavailable'
    ]
    for point in arch_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 1: Configure Master Redis Instance', bold=True)
    add_code_block(doc,
        '# On master server (redis-master):\n'
        'sudo vi /etc/redis/redis.conf\n\n'
        '# Key settings:\n'
        'bind 0.0.0.0\n'
        'port 6379\n'
        'requirepass MasterPassword123\n'
        'masterauth MasterPassword123'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 2: Configure Replica Redis Instances', bold=True)
    add_code_block(doc,
        '# On replica servers (redis-replica1, redis-replica2):\n'
        'sudo vi /etc/redis/redis.conf\n\n'
        '# Key settings:\n'
        'bind 0.0.0.0\n'
        'port 6379\n'
        'replicaof redis-master 6379\n'
        'masterauth MasterPassword123\n'
        'requirepass MasterPassword123'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Create Sentinel Configuration', bold=True)
    add_paragraph(doc, 'Purpose: Configure Sentinel to monitor the Redis master.')
    add_code_block(doc,
        '# On all sentinel nodes:\n'
        'sudo vi /etc/redis/sentinel.conf\n\n'
        '# Sentinel configuration:\n'
        'bind 0.0.0.0\n'
        'port 26379\n'
        'daemonize yes\n'
        'pidfile /var/run/redis/redis-sentinel.pid\n'
        'logfile /var/log/redis/sentinel.log\n'
        'dir /var/lib/redis\n'
        'sentinel monitor mymaster redis-master 6379 2\n'
        'sentinel auth-pass mymaster MasterPassword123\n'
        'sentinel down-after-milliseconds mymaster 5000\n'
        'sentinel parallel-syncs mymaster 1\n'
        'sentinel failover-timeout mymaster 10000'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 4: Create Sentinel Systemd Service', bold=True)
    add_code_block(doc,
        'sudo vi /etc/systemd/system/redis-sentinel.service\n\n'
        '[Unit]\n'
        'Description=Redis Sentinel\n'
        'After=network.target\n\n'
        '[Service]\n'
        'Type=forking\n'
        'User=redis\n'
        'Group=redis\n'
        'ExecStart=/usr/local/bin/redis-sentinel /etc/redis/sentinel.conf\n'
        'ExecStop=/usr/local/bin/redis-cli -p 26379 shutdown\n'
        'Restart=always\n\n'
        '[Install]\n'
        'WantedBy=multi-user.target'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 5: Start Sentinel Service', bold=True)
    add_code_block(doc,
        'sudo systemctl daemon-reload\n'
        'sudo systemctl enable redis-sentinel\n'
        'sudo systemctl start redis-sentinel\n'
        'sudo systemctl status redis-sentinel'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 6: Verify Sentinel Status', bold=True)
    add_code_block(doc,
        'redis-cli -p 26379\n'
        'SENTINEL masters\n'
        'SENTINEL replicas mymaster\n'
        'SENTINEL sentinels mymaster'
    )

    doc.add_page_break()

    # 7. Redis Cluster Configuration
    add_heading(doc, '7. Redis Cluster Configuration', 1)

    add_paragraph(doc,
        'Redis Cluster provides automatic sharding and high availability without using Sentinel. '
        'It distributes data across multiple nodes using hash slots and provides automatic failover.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Cluster Requirements:', bold=True)
    cluster_reqs = [
        'Minimum of 3 master nodes (recommended 6 nodes: 3 masters + 3 replicas)',
        'Each node requires two TCP ports: 6379 (client) and 16379 (cluster bus)',
        'All nodes must be able to communicate with each other',
        'Consistent Redis version across all nodes'
    ]
    for req in cluster_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 1: Configure Cluster Nodes', bold=True)
    add_paragraph(doc, 'Purpose: Prepare each node for cluster operation.')
    add_code_block(doc,
        '# On each node, edit redis.conf:\n'
        'sudo vi /etc/redis/redis.conf\n\n'
        '# Cluster-specific settings:\n'
        'bind 0.0.0.0\n'
        'port 6379\n'
        'cluster-enabled yes\n'
        'cluster-config-file nodes-6379.conf\n'
        'cluster-node-timeout 5000\n'
        'cluster-replica-validity-factor 10\n'
        'cluster-migration-barrier 1\n'
        'cluster-require-full-coverage yes\n'
        'appendonly yes\n'
        'requirepass ClusterPassword123\n'
        'masterauth ClusterPassword123'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 2: Start Redis on All Nodes', bold=True)
    add_code_block(doc,
        '# On each node:\n'
        'sudo systemctl restart redis'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Create Cluster', bold=True)
    add_paragraph(doc, 'Purpose: Join nodes together to form a cluster.')
    add_code_block(doc,
        '# Run on any node:\n'
        'redis-cli --cluster create \\\n'
        '  192.168.1.101:6379 \\\n'
        '  192.168.1.102:6379 \\\n'
        '  192.168.1.103:6379 \\\n'
        '  192.168.1.104:6379 \\\n'
        '  192.168.1.105:6379 \\\n'
        '  192.168.1.106:6379 \\\n'
        '  --cluster-replicas 1 \\\n'
        '  -a ClusterPassword123\n\n'
        '# Type "yes" when prompted to accept the configuration'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 4: Verify Cluster Status', bold=True)
    add_code_block(doc,
        'redis-cli -c -h 192.168.1.101 -a ClusterPassword123\n'
        'CLUSTER INFO\n'
        'CLUSTER NODES'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 5: Test Cluster Operations', bold=True)
    add_code_block(doc,
        'redis-cli -c -h 192.168.1.101 -a ClusterPassword123\n'
        'SET key1 "value1"\n'
        'GET key1\n'
        '# Note: -c flag enables cluster mode redirection'
    )

    doc.add_page_break()

    # 8. Post-Installation Configuration
    add_heading(doc, '8. Post-Installation Configuration', 1)

    add_paragraph(doc, '8.1 Memory Management', bold=True)
    add_paragraph(doc, 'Purpose: Configure memory limits and eviction policies.')
    add_code_block(doc,
        '# In redis.conf:\n'
        'maxmemory 4gb\n'
        'maxmemory-policy allkeys-lru\n\n'
        '# Available policies:\n'
        '# volatile-lru: Remove keys with expire set using LRU\n'
        '# allkeys-lru: Remove any key using LRU\n'
        '# volatile-lfu: Remove keys with expire set using LFU\n'
        '# allkeys-lfu: Remove any key using LFU\n'
        '# volatile-random: Remove random key with expire set\n'
        '# allkeys-random: Remove random key\n'
        '# volatile-ttl: Remove key with nearest expire time\n'
        '# noeviction: Return errors when memory limit is reached'
    )

    doc.add_paragraph()
    add_paragraph(doc, '8.2 Persistence Configuration', bold=True)
    add_paragraph(doc, 'Purpose: Configure data persistence strategy.')
    add_code_block(doc,
        '# RDB Snapshots (point-in-time backups):\n'
        'save 900 1      # Save after 900s if at least 1 key changed\n'
        'save 300 10     # Save after 300s if at least 10 keys changed\n'
        'save 60 10000   # Save after 60s if at least 10000 keys changed\n'
        'dbfilename dump.rdb\n\n'
        '# AOF (Append-Only File) - more durable:\n'
        'appendonly yes\n'
        'appendfilename "appendonly.aof"\n'
        'appendfsync everysec  # Options: always, everysec, no'
    )

    doc.add_paragraph()
    add_paragraph(doc, '8.3 Client Connection Limits', bold=True)
    add_code_block(doc,
        'maxclients 10000\n'
        'timeout 300  # Close idle connections after 300 seconds'
    )

    doc.add_paragraph()
    add_paragraph(doc, '8.4 Logging Configuration', bold=True)
    add_code_block(doc,
        'loglevel notice  # Options: debug, verbose, notice, warning\n'
        'logfile /var/log/redis/redis-server.log'
    )

    doc.add_page_break()

    # 9. Security Hardening
    add_heading(doc, '9. Security Hardening', 1)

    add_paragraph(doc, '9.1 Authentication', bold=True)
    add_code_block(doc,
        '# Set strong password\n'
        'requirepass YourVeryStrongPasswordHere123!\n\n'
        '# For ACL (Access Control Lists) in Redis 6+:\n'
        'aclfile /etc/redis/users.acl'
    )

    doc.add_paragraph()
    add_paragraph(doc, '9.2 Network Security', bold=True)
    add_code_block(doc,
        '# Bind to specific interfaces\n'
        'bind 127.0.0.1 192.168.1.100\n\n'
        '# Enable protected mode\n'
        'protected-mode yes'
    )

    doc.add_paragraph()
    add_paragraph(doc, '9.3 Disable Dangerous Commands', bold=True)
    add_code_block(doc,
        '# Rename or disable dangerous commands\n'
        'rename-command FLUSHDB ""\n'
        'rename-command FLUSHALL ""\n'
        'rename-command KEYS ""\n'
        'rename-command CONFIG "CONFIG_abc123"'
    )

    doc.add_paragraph()
    add_paragraph(doc, '9.4 TLS/SSL Encryption (Redis 6+)', bold=True)
    add_code_block(doc,
        'tls-port 6380\n'
        'tls-cert-file /path/to/redis.crt\n'
        'tls-key-file /path/to/redis.key\n'
        'tls-ca-cert-file /path/to/ca.crt\n'
        'tls-auth-clients no'
    )

    doc.add_page_break()

    # 10. Verification and Testing
    add_heading(doc, '10. Verification and Testing', 1)

    add_paragraph(doc, '10.1 Basic Connectivity Test', bold=True)
    add_code_block(doc,
        'redis-cli -h 127.0.0.1 -p 6379 -a YourPassword ping\n'
        '# Expected output: PONG'
    )

    doc.add_paragraph()
    add_paragraph(doc, '10.2 Information and Statistics', bold=True)
    add_code_block(doc,
        'redis-cli -a YourPassword INFO\n'
        'redis-cli -a YourPassword INFO memory\n'
        'redis-cli -a YourPassword INFO replication\n'
        'redis-cli -a YourPassword INFO stats'
    )

    doc.add_paragraph()
    add_paragraph(doc, '10.3 Performance Testing', bold=True)
    add_code_block(doc,
        '# Run benchmark\n'
        'redis-benchmark -h 127.0.0.1 -p 6379 -a YourPassword -n 100000 -c 50\n\n'
        '# Specific operations:\n'
        'redis-benchmark -h 127.0.0.1 -p 6379 -a YourPassword -t set,get -n 1000000'
    )

    doc.add_paragraph()
    add_paragraph(doc, '10.4 Data Operations Test', bold=True)
    add_code_block(doc,
        'redis-cli -a YourPassword\n'
        'SET testkey "Hello Redis 8"\n'
        'GET testkey\n'
        'EXPIRE testkey 60\n'
        'TTL testkey\n'
        'DEL testkey'
    )

    doc.add_page_break()

    # 11. Performance Tuning
    add_heading(doc, '11. Performance Tuning', 1)

    tuning_sections = [
        ('Memory Optimization', [
            'Use appropriate data structures (hashes for objects, not multiple keys)',
            'Set maxmemory and eviction policy based on use case',
            'Monitor memory fragmentation ratio',
            'Use Redis memory analyzer: redis-cli --bigkeys'
        ]),
        ('I/O Optimization', [
            'Use appendfsync everysec for balanced durability/performance',
            'Disable AOF rewrite during peak hours',
            'Use SSDs for persistence files',
            'Consider disabling persistence for cache-only workloads'
        ]),
        ('Network Optimization', [
            'Use pipelining for bulk operations',
            'Enable TCP keepalive',
            'Increase tcp-backlog for high connection rates',
            'Use Unix sockets for local connections'
        ]),
        ('CPU Optimization', [
            'Avoid KEYS command in production (use SCAN instead)',
            'Use Lua scripts to reduce round trips',
            'Monitor slow log: SLOWLOG GET 10',
            'Configure slowlog-log-slower-than appropriately'
        ])
    ]

    for category, items in tuning_sections:
        add_paragraph(doc, category + ':', bold=True)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # 12. Best Practices
    add_heading(doc, '12. Best Practices', 1)

    best_practices = [
        ('Production Deployment', [
            'Never use default configuration in production',
            'Always set a strong password with requirepass',
            'Use dedicated servers or VMs for Redis',
            'Implement monitoring and alerting',
            'Regular backups of RDB/AOF files',
            'Test failover scenarios regularly'
        ]),
        ('High Availability', [
            'Use Redis Sentinel for automatic failover',
            'Deploy at least 3 Sentinel instances',
            'Use Redis Cluster for horizontal scaling',
            'Configure appropriate replication settings',
            'Monitor replication lag'
        ]),
        ('Monitoring', [
            'Monitor memory usage and fragmentation',
            'Track connected clients and blocked clients',
            'Monitor keyspace hits/misses ratio',
            'Set up alerts for high memory usage',
            'Use Redis INFO command for metrics',
            'Consider Prometheus exporter for metrics'
        ]),
        ('Backup and Recovery', [
            'Schedule regular RDB snapshots',
            'Store backups on separate storage',
            'Test restore procedures regularly',
            'Use AOF for critical data',
            'Implement automated backup scripts'
        ])
    ]

    for category, items in best_practices:
        add_paragraph(doc, category + ':', bold=True)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # 13. Troubleshooting
    add_heading(doc, '13. Troubleshooting', 1)

    troubleshooting = [
        ['Issue', 'Possible Cause', 'Solution'],
        ['High memory usage', 'Memory leak or no eviction policy', 'Check maxmemory settings, analyze with --bigkeys'],
        ['Slow queries', 'Blocking commands or large datasets', 'Check SLOWLOG, avoid KEYS command'],
        ['Connection refused', 'Firewall or bind address', 'Verify firewall rules and bind configuration'],
        ['Replication lag', 'Network issues or heavy load', 'Check network, reduce load on master'],
        ['AOF rewrite taking too long', 'Large dataset or slow disk', 'Use SSDs, tune no-appendfsync-on-rewrite']
    ]

    add_table_with_header(doc, troubleshooting[0], troubleshooting[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Useful Diagnostic Commands:', bold=True)
    add_code_block(doc,
        '# Check memory usage\n'
        'redis-cli -a YourPassword INFO memory\n\n'
        '# Find large keys\n'
        'redis-cli -a YourPassword --bigkeys\n\n'
        '# Check slow queries\n'
        'redis-cli -a YourPassword SLOWLOG GET 10\n\n'
        '# Monitor commands in real-time\n'
        'redis-cli -a YourPassword MONITOR\n\n'
        '# Check client connections\n'
        'redis-cli -a YourPassword CLIENT LIST\n\n'
        '# Check replication status\n'
        'redis-cli -a YourPassword INFO replication'
    )

    doc.add_page_break()

    # Footer
    add_heading(doc, 'Document Information', 1)
    doc_info = [
        ['Field', 'Value'],
        ['Document Version', '1.0'],
        ['Last Updated', '2026-01-14'],
        ['Platform', 'Red Hat Enterprise Linux 8'],
        ['Software Version', 'Redis 8.x'],
        ['Purpose', 'Installation and Configuration Guide']
    ]
    add_table_with_header(doc, doc_info[0], doc_info[1:])

    return doc

if __name__ == '__main__':
    doc = create_redis_document()
    doc.save('Redis_8.x_RHEL8_Installation_Guide.docx')
    print("Redis documentation created successfully!")
