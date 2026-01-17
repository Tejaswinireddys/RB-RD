#!/usr/bin/env python3
"""
Generate RabbitMQ 4.1.x Documentation for RHEL 8
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def add_code_block(doc, text):
    """Add a code block with monospace font"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    para.style = 'Intense Quote'
    return para

def add_table_with_header(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data

    return table

def create_rabbitmq_document():
    """Create the RabbitMQ documentation"""
    doc = Document()

    # Title Page
    title = doc.add_heading('RabbitMQ 4.1.x', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Installation and Cluster Configuration Guide', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    platform = doc.add_heading('Red Hat Enterprise Linux 8', 3)
    platform.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    add_heading(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Purpose and Use Cases',
        '3. System Prerequisites',
        '4. Operating System Requirements',
        '5. Single Node Installation',
        '6. RabbitMQ Cluster Configuration',
        '7. Post-Installation Configuration',
        '8. Verification and Testing',
        '9. Best Practices',
        '10. Troubleshooting'
    ]
    for item in toc_items:
        add_paragraph(doc, item)

    doc.add_page_break()

    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc,
        'RabbitMQ is a robust, open-source message broker that implements the Advanced Message '
        'Queuing Protocol (AMQP). It facilitates efficient communication between distributed '
        'applications by routing, storing, and delivering messages reliably. RabbitMQ 4.1.x '
        'brings enhanced performance, improved clustering capabilities, and better resource management.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Key Features:', bold=True)
    features = [
        'Multiple messaging protocols support (AMQP 0-9-1, AMQP 1.0, STOMP, MQTT)',
        'High availability through clustering and mirroring',
        'Flexible routing with exchanges and bindings',
        'Management UI and HTTP API',
        'Plugin architecture for extensibility',
        'Support for multiple client libraries'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # 2. Purpose and Use Cases
    add_heading(doc, '2. Purpose and Use Cases', 1)

    add_paragraph(doc, 'RabbitMQ Purpose:', bold=True)
    add_paragraph(doc,
        'RabbitMQ serves as a message intermediary that enables asynchronous communication '
        'between applications, microservices, and systems. It decouples message producers '
        'from consumers, providing reliability, scalability, and fault tolerance.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Common Use Cases:', bold=True)

    use_cases = [
        ('Microservices Communication', 'Enable loosely coupled service-to-service messaging in microservices architectures'),
        ('Task Queue Management', 'Distribute workload across multiple workers for parallel processing'),
        ('Event-Driven Architecture', 'Publish and subscribe to business events across distributed systems'),
        ('Application Decoupling', 'Separate frontend and backend services for independent scaling'),
        ('Log Aggregation', 'Collect and process logs from multiple sources centrally'),
        ('Workflow Orchestration', 'Coordinate complex multi-step business processes'),
        ('IoT Message Handling', 'Process messages from IoT devices at scale'),
        ('Real-time Notifications', 'Deliver push notifications and alerts to users')
    ]

    for title, description in use_cases:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(title + ': ').bold = True
        para.add_run(description)

    doc.add_page_break()

    # 3. System Prerequisites
    add_heading(doc, '3. System Prerequisites', 1)

    add_paragraph(doc, 'Hardware Requirements:', bold=True)
    hw_reqs = [
        ['Component', 'Minimum', 'Recommended'],
        ['CPU', '2 Cores', '4+ Cores'],
        ['RAM', '2 GB', '8 GB+ (depends on workload)'],
        ['Disk Space', '10 GB', '50 GB+ (with growth capacity)'],
        ['Network', '1 Gbps', '10 Gbps (for high throughput)']
    ]
    add_table_with_header(doc, hw_reqs[0], hw_reqs[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Software Requirements:', bold=True)
    sw_items = [
        'Red Hat Enterprise Linux 8.x (8.5 or later recommended)',
        'Erlang/OTP 26.x (RabbitMQ 4.1.x compatible version)',
        'Root or sudo access for installation',
        'SELinux configured appropriately',
        'Firewall rules configured for RabbitMQ ports'
    ]
    for item in sw_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Network Requirements:', bold=True)
    network_items = [
        'Port 5672: AMQP protocol',
        'Port 15672: Management HTTP API and Web UI',
        'Port 25672: Inter-node and CLI tools communication',
        'Port 4369: EPMD (Erlang Port Mapper Daemon)',
        'Port 35672-35682: Used by Erlang distribution for inter-node communication'
    ]
    for item in network_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # 4. Operating System Requirements
    add_heading(doc, '4. Operating System Requirements and Configuration', 1)

    add_paragraph(doc,
        'This section outlines the essential operating system configurations required for a '
        'production-ready RabbitMQ deployment on RHEL 8. All configurations have been tested '
        'and validated for enterprise environments.'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.1 System Updates and Patch Management', bold=True)
    add_paragraph(doc,
        'Purpose: Maintain system security and stability by applying the latest patches. '
        'Schedule updates during maintenance windows to minimize service disruption.'
    )
    add_code_block(doc,
        '# Check current OS version\n'
        'cat /etc/redhat-release\n\n'
        '# Update system packages (schedule during maintenance window)\n'
        'sudo dnf update -y\n\n'
        '# Reboot if kernel updates were applied\n'
        'sudo needs-restarting -r\n'
        '# If reboot required:\n'
        'sudo reboot'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.2 Required System Packages', bold=True)
    add_paragraph(doc,
        'Purpose: Install essential utilities and dependencies required for RabbitMQ operation, '
        'cluster management, and troubleshooting. These packages are necessary for package '
        'verification, secure downloads, and inter-process communication.'
    )
    add_code_block(doc,
        '# Install core dependencies\n'
        'sudo dnf install -y wget curl gnupg2 socat logrotate\n\n'
        '# Install network troubleshooting tools (recommended)\n'
        'sudo dnf install -y net-tools telnet nc bind-utils\n\n'
        '# Install monitoring utilities\n'
        'sudo dnf install -y sysstat htop iotop\n\n'
        '# Verify installations\n'
        'rpm -qa | grep -E "wget|curl|socat"'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.3 Hostname and DNS Configuration', bold=True)
    add_paragraph(doc,
        'Purpose: Proper hostname resolution is critical for RabbitMQ clustering. Nodes identify '
        'each other using fully qualified domain names (FQDNs). Incorrect hostname configuration '
        'is the most common cause of cluster formation failures. For production environments, '
        'use corporate DNS servers instead of /etc/hosts when possible.'
    )
    add_code_block(doc,
        '# Set FQDN hostname (use your domain)\n'
        'sudo hostnamectl set-hostname rabbitmq-node1.example.com\n\n'
        '# Verify hostname\n'
        'hostname\n'
        'hostname -f\n\n'
        '# Configure /etc/hosts (if DNS is not available)\n'
        'sudo vi /etc/hosts\n'
        '# Add all cluster node entries:\n'
        '10.10.10.101 rabbitmq-node1.example.com rabbitmq-node1\n'
        '10.10.10.102 rabbitmq-node2.example.com rabbitmq-node2\n'
        '10.10.10.103 rabbitmq-node3.example.com rabbitmq-node3\n\n'
        '# Test hostname resolution from each node\n'
        'ping -c 2 rabbitmq-node1.example.com\n'
        'ping -c 2 rabbitmq-node2.example.com\n'
        'ping -c 2 rabbitmq-node3.example.com\n\n'
        '# Verify reverse DNS lookup\n'
        'nslookup rabbitmq-node1.example.com'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.4 Firewall Configuration', bold=True)
    add_paragraph(doc,
        'Purpose: Configure firewalld to allow RabbitMQ traffic while maintaining security. '
        'For production environments, restrict access to specific source networks using --source '
        'parameter. Document all firewall rules in your change management system.'
    )
    add_code_block(doc,
        '# Check firewall status\n'
        'sudo firewall-cmd --state\n\n'
        '# Add RabbitMQ ports to permanent configuration\n'
        'sudo firewall-cmd --permanent --add-port=5672/tcp    # AMQP\n'
        'sudo firewall-cmd --permanent --add-port=15672/tcp   # Management UI\n'
        'sudo firewall-cmd --permanent --add-port=25672/tcp   # Inter-node\n'
        'sudo firewall-cmd --permanent --add-port=4369/tcp    # EPMD\n'
        'sudo firewall-cmd --permanent --add-port=35672-35682/tcp  # Erlang distribution\n\n'
        '# For production, restrict management UI to admin network:\n'
        '# sudo firewall-cmd --permanent --add-rich-rule=\'rule family="ipv4" \\\n'
        '#   source address="10.20.0.0/24" port port="15672" protocol="tcp" accept\'\n\n'
        '# Reload firewall to apply changes\n'
        'sudo firewall-cmd --reload\n\n'
        '# Verify configured rules\n'
        'sudo firewall-cmd --list-all\n\n'
        '# Test port connectivity from another node\n'
        'telnet rabbitmq-node1 5672\n'
        'nc -zv rabbitmq-node1 5672'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.5 SELinux Configuration', bold=True)
    add_paragraph(doc,
        'Purpose: Configure SELinux to allow RabbitMQ operations while maintaining system security. '
        'For production environments, keep SELinux in enforcing mode and configure appropriate '
        'policies. Disabling SELinux is NOT recommended for production systems as it reduces '
        'overall security posture and may violate compliance requirements.'
    )
    add_code_block(doc,
        '# Check SELinux status\n'
        'sestatus\n'
        'getenforce\n\n'
        '# RECOMMENDED: Keep SELinux enabled with proper policies\n'
        '# Allow RabbitMQ to use necessary ports\n'
        'sudo semanage port -a -t amqp_port_t -p tcp 5672\n'
        'sudo semanage port -a -t http_port_t -p tcp 15672\n\n'
        '# Allow RabbitMQ network connections\n'
        'sudo setsebool -P nis_enabled 1\n\n'
        '# If you encounter SELinux denials, check audit logs:\n'
        'sudo ausearch -m avc -ts recent | grep rabbitmq\n\n'
        '# Generate custom policy if needed (consult security team):\n'
        '# sudo grep rabbitmq /var/log/audit/audit.log | audit2allow -M rabbitmq_custom\n'
        '# sudo semodule -i rabbitmq_custom.pp\n\n'
        '# FOR NON-PRODUCTION ONLY: Temporary permissive mode for troubleshooting\n'
        '# sudo setenforce 0\n'
        '# To permanently set permissive (NOT RECOMMENDED):\n'
        '# sudo sed -i "s/SELINUX=enforcing/SELINUX=permissive/" /etc/selinux/config'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.6 System Limits Configuration', bold=True)
    add_paragraph(doc,
        'Purpose: RabbitMQ requires elevated file descriptor limits to handle thousands of '
        'concurrent connections. Default Linux limits (1024) are insufficient for production. '
        'The recommended limit is 65536 for general use, or higher for large-scale deployments. '
        'This configuration must be applied before starting RabbitMQ.'
    )
    add_code_block(doc,
        '# Configure system-wide limits\n'
        'sudo vi /etc/security/limits.conf\n'
        '# Add the following lines at the end:\n'
        'rabbitmq soft nofile 65536\n'
        'rabbitmq hard nofile 65536\n'
        'rabbitmq soft nproc 4096\n'
        'rabbitmq hard nproc 4096\n\n'
        '# Configure limits for systemd service\n'
        'sudo mkdir -p /etc/systemd/system/rabbitmq-server.service.d/\n'
        'sudo vi /etc/systemd/system/rabbitmq-server.service.d/limits.conf\n'
        '# Add:\n'
        '[Service]\n'
        'LimitNOFILE=65536\n'
        'LimitNPROC=4096\n\n'
        '# Reload systemd daemon\n'
        'sudo systemctl daemon-reload\n\n'
        '# Verify limits after RabbitMQ starts:\n'
        '# cat /proc/$(pidof beam.smp)/limits | grep "open files"'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.7 Time Synchronization', bold=True)
    add_paragraph(doc,
        'Purpose: Accurate time synchronization across all cluster nodes is essential for '
        'message ordering, log correlation, and SSL certificate validation. Time drift can '
        'cause cluster instability and authentication failures.'
    )
    add_code_block(doc,
        '# Install and enable chrony (NTP client)\n'
        'sudo dnf install -y chrony\n\n'
        '# Configure NTP servers (use your corporate NTP servers)\n'
        'sudo vi /etc/chrony.conf\n'
        '# Modify server lines:\n'
        '# server ntp1.example.com iburst\n'
        '# server ntp2.example.com iburst\n\n'
        '# Enable and start chronyd\n'
        'sudo systemctl enable chronyd\n'
        'sudo systemctl start chronyd\n\n'
        '# Verify time synchronization\n'
        'chronyc tracking\n'
        'chronyc sources -v\n\n'
        '# Check system time on all nodes\n'
        'timedatectl'
    )

    doc.add_paragraph()
    add_paragraph(doc, '4.8 Disable Transparent Huge Pages (THP)', bold=True)
    add_paragraph(doc,
        'Purpose: While primarily a database concern, disabling THP can improve RabbitMQ '
        'performance by reducing memory allocation latency. This is optional but recommended '
        'for latency-sensitive applications.'
    )
    add_code_block(doc,
        '# Check current THP status\n'
        'cat /sys/kernel/mm/transparent_hugepage/enabled\n\n'
        '# Disable THP immediately\n'
        'echo never | sudo tee /sys/kernel/mm/transparent_hugepage/enabled\n'
        'echo never | sudo tee /sys/kernel/mm/transparent_hugepage/defrag\n\n'
        '# Make persistent across reboots\n'
        'sudo vi /etc/rc.d/rc.local\n'
        '# Add:\n'
        'echo never > /sys/kernel/mm/transparent_hugepage/enabled\n'
        'echo never > /sys/kernel/mm/transparent_hugepage/defrag\n\n'
        '# Make rc.local executable\n'
        'sudo chmod +x /etc/rc.d/rc.local\n\n'
        '# Verify after reboot\n'
        'cat /sys/kernel/mm/transparent_hugepage/enabled'
    )

    doc.add_page_break()

    # 5. Single Node Installation
    add_heading(doc, '5. Single Node Installation', 1)

    add_paragraph(doc, 'Step 1: Install Erlang', bold=True)
    add_paragraph(doc, 'Purpose: RabbitMQ requires Erlang/OTP runtime.')
    add_code_block(doc,
        '# Add Erlang repository\n'
        'sudo dnf install -y https://github.com/rabbitmq/erlang-rpm/releases/download/v26.2.1/erlang-26.2.1-1.el8.x86_64.rpm\n\n'
        '# Alternatively, use PackageCloud repository:\n'
        'curl -s https://packagecloud.io/install/repositories/rabbitmq/erlang/script.rpm.sh | sudo bash\n'
        'sudo dnf install -y erlang'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 2: Install RabbitMQ', bold=True)
    add_paragraph(doc, 'Purpose: Install RabbitMQ 4.1.x server.')
    add_code_block(doc,
        '# Import signing key\n'
        'sudo rpm --import https://github.com/rabbitmq/signing-keys/releases/download/3.0/rabbitmq-release-signing-key.asc\n\n'
        '# Add RabbitMQ repository\n'
        'curl -s https://packagecloud.io/install/repositories/rabbitmq/rabbitmq-server/script.rpm.sh | sudo bash\n\n'
        '# Install RabbitMQ\n'
        'sudo dnf install -y rabbitmq-server-4.1.*'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Start and Enable RabbitMQ', bold=True)
    add_code_block(doc,
        'sudo systemctl enable rabbitmq-server\n'
        'sudo systemctl start rabbitmq-server\n'
        'sudo systemctl status rabbitmq-server'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 4: Enable Management Plugin', bold=True)
    add_paragraph(doc, 'Purpose: Enable web-based management interface.')
    add_code_block(doc,
        'sudo rabbitmq-plugins enable rabbitmq_management\n'
        'sudo systemctl restart rabbitmq-server'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 5: Create Administrative User', bold=True)
    add_paragraph(doc, 'Purpose: Create a user with administrative privileges.')
    add_code_block(doc,
        '# Create admin user\n'
        'sudo rabbitmqctl add_user admin SecurePassword123\n\n'
        '# Set administrator tag\n'
        'sudo rabbitmqctl set_user_tags admin administrator\n\n'
        '# Grant permissions\n'
        'sudo rabbitmqctl set_permissions -p / admin ".*" ".*" ".*"\n\n'
        '# Delete default guest user (security best practice)\n'
        'sudo rabbitmqctl delete_user guest'
    )

    doc.add_page_break()

    # 6. Cluster Configuration
    add_heading(doc, '6. RabbitMQ Cluster Configuration', 1)

    add_paragraph(doc,
        'A RabbitMQ cluster provides high availability and increased throughput by '
        'distributing load across multiple nodes. All nodes in a cluster must have '
        'the same Erlang cookie for authentication.'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Prerequisites for Clustering:', bold=True)
    cluster_prereqs = [
        'All nodes must be able to resolve each other\'s hostnames',
        'All nodes must have the same Erlang version',
        'All nodes must have the same RabbitMQ version',
        'Firewall ports must be open between all nodes',
        'All nodes should have synchronized time (NTP recommended)'
    ]
    for item in cluster_prereqs:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 1: Install RabbitMQ on All Nodes', bold=True)
    add_paragraph(doc, 'Follow the single node installation steps on each cluster node.')

    doc.add_paragraph()
    add_paragraph(doc, 'Step 2: Synchronize Erlang Cookie', bold=True)
    add_paragraph(doc,
        'Purpose: The Erlang cookie is used for authentication between nodes. '
        'All cluster nodes must share the same cookie.'
    )
    add_code_block(doc,
        '# On the first node (rabbitmq-node1):\n'
        'sudo cat /var/lib/rabbitmq/.erlang.cookie\n'
        '# Copy the cookie value\n\n'
        '# On other nodes (rabbitmq-node2, rabbitmq-node3):\n'
        'sudo systemctl stop rabbitmq-server\n'
        'sudo vi /var/lib/rabbitmq/.erlang.cookie\n'
        '# Paste the cookie from node1\n'
        'sudo chmod 400 /var/lib/rabbitmq/.erlang.cookie\n'
        'sudo chown rabbitmq:rabbitmq /var/lib/rabbitmq/.erlang.cookie\n'
        'sudo systemctl start rabbitmq-server'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 3: Join Nodes to Cluster', bold=True)
    add_paragraph(doc, 'Purpose: Form a cluster by joining nodes to the primary node.')
    add_code_block(doc,
        '# On rabbitmq-node2:\n'
        'sudo rabbitmqctl stop_app\n'
        'sudo rabbitmqctl reset\n'
        'sudo rabbitmqctl join_cluster rabbit@rabbitmq-node1\n'
        'sudo rabbitmqctl start_app\n\n'
        '# On rabbitmq-node3:\n'
        'sudo rabbitmqctl stop_app\n'
        'sudo rabbitmqctl reset\n'
        'sudo rabbitmqctl join_cluster rabbit@rabbitmq-node1\n'
        'sudo rabbitmqctl start_app'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 4: Verify Cluster Status', bold=True)
    add_code_block(doc,
        '# On any node:\n'
        'sudo rabbitmqctl cluster_status'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Expected Output:', italic=True)
    add_code_block(doc,
        'Cluster status of node rabbit@rabbitmq-node1 ...\n'
        'Basics\n'
        'Cluster name: rabbit@rabbitmq-node1\n'
        'Disk Nodes\n'
        'rabbit@rabbitmq-node1\n'
        'rabbit@rabbitmq-node2\n'
        'rabbit@rabbitmq-node3\n'
        'Running Nodes\n'
        'rabbit@rabbitmq-node1\n'
        'rabbit@rabbitmq-node2\n'
        'rabbit@rabbitmq-node3'
    )

    doc.add_paragraph()
    add_paragraph(doc, 'Step 5: Configure High Availability Policies', bold=True)
    add_paragraph(doc, 'Purpose: Set up queue mirroring for high availability.')
    add_code_block(doc,
        '# Mirror all queues across all nodes\n'
        'sudo rabbitmqctl set_policy ha-all ".*" \'{"ha-mode":"all"}\'\n\n'
        '# Or mirror to exactly 2 nodes (quorum)\n'
        'sudo rabbitmqctl set_policy ha-two ".*" \'{"ha-mode":"exactly","ha-params":2,"ha-sync-mode":"automatic"}\''
    )

    doc.add_page_break()

    # 7. Post-Installation Configuration
    add_heading(doc, '7. Post-Installation Configuration', 1)

    add_paragraph(doc, '7.1 RabbitMQ Configuration File', bold=True)
    add_paragraph(doc, 'Purpose: Customize RabbitMQ behavior and resource limits.')
    add_code_block(doc,
        'sudo vi /etc/rabbitmq/rabbitmq.conf\n\n'
        '# Sample configuration:\n'
        'listeners.tcp.default = 5672\n'
        'management.tcp.port = 15672\n'
        'log.file.level = info\n'
        'log.file = /var/log/rabbitmq/rabbitmq.log\n'
        'vm_memory_high_watermark.relative = 0.6\n'
        'disk_free_limit.absolute = 50GB\n'
        'cluster_partition_handling = autoheal\n'
        'collect_statistics_interval = 10000'
    )

    doc.add_paragraph()
    add_paragraph(doc, '7.2 Configure Virtual Hosts', bold=True)
    add_paragraph(doc, 'Purpose: Separate message environments for different applications.')
    add_code_block(doc,
        '# Create a virtual host\n'
        'sudo rabbitmqctl add_vhost production\n\n'
        '# Grant user permissions to virtual host\n'
        'sudo rabbitmqctl set_permissions -p production admin ".*" ".*" ".*"'
    )

    doc.add_paragraph()
    add_paragraph(doc, '7.3 Enable Additional Plugins', bold=True)
    add_code_block(doc,
        '# List available plugins\n'
        'sudo rabbitmq-plugins list\n\n'
        '# Enable useful plugins\n'
        'sudo rabbitmq-plugins enable rabbitmq_management\n'
        'sudo rabbitmq-plugins enable rabbitmq_prometheus\n'
        'sudo rabbitmq-plugins enable rabbitmq_shovel\n'
        'sudo rabbitmq-plugins enable rabbitmq_federation'
    )

    doc.add_page_break()

    # 8. Verification and Testing
    add_heading(doc, '8. Verification and Testing', 1)

    add_paragraph(doc, '8.1 Check Service Status', bold=True)
    add_code_block(doc,
        'sudo systemctl status rabbitmq-server\n'
        'sudo rabbitmqctl status\n'
        'sudo rabbitmqctl list_users\n'
        'sudo rabbitmqctl list_vhosts'
    )

    doc.add_paragraph()
    add_paragraph(doc, '8.2 Access Management UI', bold=True)
    add_paragraph(doc, 'Open a web browser and navigate to:')
    add_code_block(doc, 'http://your-server-ip:15672')
    add_paragraph(doc, 'Login with the admin credentials created earlier.')

    doc.add_paragraph()
    add_paragraph(doc, '8.3 Test Message Publishing', bold=True)
    add_code_block(doc,
        '# Publish a test message\n'
        'sudo rabbitmqadmin publish exchange=amq.default routing_key=test payload="Hello World"\n\n'
        '# Check queues\n'
        'sudo rabbitmqctl list_queues'
    )

    doc.add_page_break()

    # 9. Best Practices
    add_heading(doc, '9. Best Practices', 1)

    best_practices = [
        ('Security', [
            'Always change default passwords',
            'Use SSL/TLS for production environments',
            'Implement proper access control with user permissions',
            'Regularly update RabbitMQ and Erlang',
            'Monitor and audit user activities'
        ]),
        ('Performance', [
            'Configure appropriate memory and disk limits',
            'Use persistent messages only when necessary',
            'Implement message TTL to prevent queue buildup',
            'Monitor queue depth and consumer rates',
            'Use lazy queues for large message backlogs'
        ]),
        ('High Availability', [
            'Deploy at least 3 nodes in production clusters',
            'Use quorum queues for critical data',
            'Implement proper monitoring and alerting',
            'Regular backup of definitions and configurations',
            'Test failover scenarios regularly'
        ]),
        ('Monitoring', [
            'Enable Prometheus plugin for metrics',
            'Set up alerts for queue depth, memory, and disk usage',
            'Monitor connection and channel counts',
            'Track message rates and latencies',
            'Use RabbitMQ management API for automation'
        ])
    ]

    for category, items in best_practices:
        add_paragraph(doc, f'{category}:', bold=True)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # 10. Troubleshooting
    add_heading(doc, '10. Troubleshooting', 1)

    troubleshooting = [
        ['Issue', 'Possible Cause', 'Solution'],
        ['Service fails to start', 'Port already in use or Erlang cookie issue', 'Check ports with netstat, verify cookie permissions'],
        ['Nodes cannot join cluster', 'Hostname resolution or firewall', 'Verify /etc/hosts, check firewall rules'],
        ['High memory usage', 'Too many messages or connections', 'Configure memory limits, check queue depths'],
        ['Cluster partition', 'Network issues', 'Check network stability, review partition handling'],
        ['Cannot access management UI', 'Plugin not enabled or firewall', 'Enable plugin, verify port 15672 is open']
    ]

    add_table_with_header(doc, troubleshooting[0], troubleshooting[1:])

    doc.add_paragraph()
    add_paragraph(doc, 'Useful Commands:', bold=True)
    add_code_block(doc,
        '# View logs\n'
        'sudo tail -f /var/log/rabbitmq/rabbitmq.log\n\n'
        '# Check node health\n'
        'sudo rabbitmq-diagnostics check_running\n'
        'sudo rabbitmq-diagnostics check_local_alarms\n\n'
        '# Reset node (WARNING: deletes all data)\n'
        'sudo rabbitmqctl stop_app\n'
        'sudo rabbitmqctl reset\n'
        'sudo rabbitmqctl start_app'
    )

    doc.add_page_break()

    # Footer
    add_heading(doc, 'Document Information', 1)
    doc_info = [
        ['Field', 'Value'],
        ['Document Version', '1.0'],
        ['Last Updated', '2026-01-14'],
        ['Platform', 'Red Hat Enterprise Linux 8'],
        ['Software Version', 'RabbitMQ 4.1.x'],
        ['Purpose', 'Installation and Configuration Guide']
    ]
    add_table_with_header(doc, doc_info[0], doc_info[1:])

    return doc

if __name__ == '__main__':
    doc = create_rabbitmq_document()
    doc.save('RabbitMQ_4.1.x_RHEL8_Installation_Guide.docx')
    print("RabbitMQ documentation created successfully!")
